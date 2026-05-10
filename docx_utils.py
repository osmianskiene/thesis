from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

C_BLACK = (0,   0,   0)
C_GREEN = (0,   130, 0)
C_RED   = (192, 0,   0)
C_GRAY  = (64,  64,  64)

FONT_BODY     = 11
FONT_SUBGROUP =  7
FONT_LEGEND   =  8


def dir_color(value: float, direction: str | None) -> tuple:
    if direction == 'down':
        return C_GREEN if value < 0 else (C_RED if value > 0 else C_GRAY)
    if direction == 'up':
        return C_GREEN if value > 0 else (C_RED if value < 0 else C_GRAY)
    return C_GRAY


def add_run(para: Paragraph, text: str, size: float, bold: bool = False, rgb: tuple = C_BLACK) -> Run:
    r = para.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor(*rgb)
    return r


def clear_cell(cell: _Cell) -> None:
    # Reset to a single empty paragraph, preserving paragraph formatting properties.
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for child in list(p0._element):
        if child.tag != qn('w:pPr'):
            p0._element.remove(child)


def new_para(cell: _Cell) -> Paragraph:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def set_col_widths(table: Table, widths_cm: list) -> None:
    # python-docx doesn't expose column widths directly, so we manipulate the
    # underlying OOXML tree. Column widths in OOXML are stored in twips
    # (twentieths of a point). python-docx's Cm() returns EMUs (English Metric
    # Units), and 635 EMU = 1 twip.
    twips = [int(Cm(w) / 635) for w in widths_cm]  # convert cm → twips
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    for gc in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(gc)
    for tw in twips:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(tw))
        tblGrid.append(gc)
    for row in table.rows:
        for cell, tw in zip(row.cells, twips):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(tw))
            tcW.set(qn('w:type'), 'dxa')
