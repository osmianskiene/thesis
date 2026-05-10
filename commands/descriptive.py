from pathlib import Path
from typing import cast

from docx import Document
from docx.shared import Pt, Cm
from docx.styles.style import _ParagraphStyle
from docx.table import _Cell

from biomarkers import BIOMARKERS
from data import load_data
from docx_utils import (
    C_BLACK, C_GRAY, FONT_BODY, FONT_LEGEND, FONT_SUBGROUP,
    add_run, clear_cell, dir_color, new_para, set_col_widths,
)
from stats import central, fmt_stat, get_subgroups


def _subgroup_line(cell: _Cell, groups: list, is_delta: bool, direction: str | None) -> None:
    para = new_para(cell)
    for i, (label, series) in enumerate(groups):
        if i:
            add_run(para, ', ', FONT_SUBGROUP, rgb=C_GRAY)
        add_run(para, label, FONT_SUBGROUP, bold=True)
        color = dir_color(central(series), direction) if is_delta else C_GRAY
        add_run(para, ' ' + fmt_stat(series), FONT_SUBGROUP, rgb=color)


def _fill_stat_cell(cell: _Cell, subgroups: dict, is_delta: bool, direction: str | None = None) -> None:
    clear_cell(cell)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(0)
    color = dir_color(central(subgroups['all']), direction) if is_delta else C_BLACK
    add_run(para, fmt_stat(subgroups['all']), FONT_BODY, rgb=color)
    _subgroup_line(cell, [('T', subgroups['T']), ('C', subgroups['C'])], is_delta, direction)
    _subgroup_line(cell, [('M', subgroups['M']), ('F', subgroups['F'])], is_delta, direction)
    _subgroup_line(cell, [('Y', subgroups['Y']), ('A', subgroups['A']), ('O', subgroups['O'])], is_delta, direction)


def _fill_count_cell(cell: _Cell, subgroups: dict) -> None:
    clear_cell(cell)
    total = subgroups['all']
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, f'{total} (100%)', FONT_BODY)

    def count_line(groups: list) -> None:
        p = new_para(cell)
        for i, (label, n) in enumerate(groups):
            if i:
                add_run(p, ', ', FONT_SUBGROUP, rgb=C_GRAY)
            add_run(p, label, FONT_SUBGROUP, bold=True)
            add_run(p, f' {n} ({n / total * 100:.1f}%)', FONT_SUBGROUP, rgb=C_GRAY)

    count_line([('T', subgroups['T']), ('C', subgroups['C'])])
    count_line([('M', subgroups['M']), ('F', subgroups['F'])])
    count_line([('Y', subgroups['Y']), ('A', subgroups['A']), ('O', subgroups['O'])])


def _fill_name_cell(cell: _Cell, name: str, bold: bool = False) -> None:
    clear_cell(cell)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, name, FONT_BODY, bold=bold)


def descriptive_docx() -> None:
    df = load_data()
    Path('output').mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    name_cell, t0_cell, delta_cell = table.rows[0].cells
    _fill_name_cell(name_cell, 'Biomarker', bold=True)
    _fill_name_cell(t0_cell, 'T0', bold=True)
    _fill_name_cell(delta_cell, 'Δ', bold=True)

    participant_subgroups = {
        'all': len(df),
        'T': int((df['Treatment'] == 1).sum()),
        'C': int((df['Treatment'] == 0).sum()),
        'M': int((df['Sex'] == 'm').sum()),
        'F': int((df['Sex'] == 'f').sum()),
        'Y': int((df['Age Group'] == 'Young').sum()),
        'A': int((df['Age Group'] == 'Average').sum()),
        'O': int((df['Age Group'] == 'Old').sum()),
    }
    name_cell, t0_cell, delta_cell = table.add_row().cells
    _fill_name_cell(name_cell, 'Participants')
    _fill_count_cell(t0_cell, participant_subgroups)
    _fill_count_cell(delta_cell, participant_subgroups)

    name_cell, t0_cell, delta_cell = table.add_row().cells
    _fill_name_cell(name_cell, 'Chronological Age')
    _fill_stat_cell(t0_cell, get_subgroups(df, 'Chronological Age T0'), is_delta=False)
    _fill_stat_cell(delta_cell, get_subgroups(df, 'Chronological Age Δ'), is_delta=True)

    for bm in BIOMARKERS:
        name_cell, t0_cell, delta_cell = table.add_row().cells
        _fill_name_cell(name_cell, bm.name)
        _fill_stat_cell(t0_cell, get_subgroups(df, bm.t0_col), is_delta=False)
        _fill_stat_cell(delta_cell, get_subgroups(df, bm.delta_col), is_delta=True, direction=bm.direction)

    set_col_widths(table, [4.5, 6.5, 6.5])

    doc.add_paragraph()
    legend = [
        ('T0',         ': baseline measurement.'),
        ('Δ',          ': change from baseline to follow-up (T1−T0).'),
        ('Statistics', ': mean±SD (Shapiro–Wilk p > 0.05); median [Q1–Q3] otherwise.'),
        ('Subgroups',  ': T = treatment, C = control; M = male, F = female; '
                       'Y = Young (<40 yr), A = Average (40–54 yr), O = Old (>54 yr).'),
        ('Δ color',    ': green = improved; red = worsened. '
                       'Higher is better: NAD⁺, HDL. Lower is better: all others.'),
    ]
    for key, val in legend:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, key, FONT_LEGEND, bold=True)
        add_run(para, val, FONT_LEGEND)

    out = 'output/a1_descriptive.docx'
    doc.save(out)
    print(f'Saved: {out}')
