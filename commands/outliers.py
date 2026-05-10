from pathlib import Path
from typing import cast

import numpy as np
import statsmodels.api as sm
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from statsmodels.stats.outliers_influence import OLSInfluence

from biomarkers import BIOMARKERS
from data import load_data
from docx_utils import C_BLACK, C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths


# ── shared helpers ────────────────────────────────────────────────────────────

def _make_doc() -> Document:
    """Create a blank Word document with standard margins and Arial font."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)
    return doc


def _header_row(table, headers: list[str]) -> None:
    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        add_run(para, text, FONT_BODY, bold=True)


def _data_row(table, values: list[str]) -> None:
    for cell, val in zip(table.add_row().cells, values):
        clear_cell(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        add_run(para, val, FONT_BODY, rgb=C_BLACK)


def _add_legend(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, text, FONT_LEGEND, rgb=C_GRAY)


# ── data collection ───────────────────────────────────────────────────────────

def _collect_zscore_flags(df) -> list[dict]:
    """Return rows with |z| > 3 on each biomarker's Δ column."""
    flags = []
    for bm in BIOMARKERS:
        col = bm.delta_col
        if col not in df.columns:
            continue
        series = df[col].dropna()
        z = (series - series.mean()) / series.std()
        for idx, zval in z[np.abs(z) > 3].items():
            flags.append({
                'biomarker': bm.name,
                'row':       df.loc[idx, 'Row'],
                'group':     df.loc[idx, 'Group'],
                'delta':     df.loc[idx, col],
                'z':         zval,
            })
    flags.sort(key=lambda r: abs(r['z']), reverse=True)
    return flags


def _collect_ancova_flags(df) -> tuple[list[dict], list[dict]]:
    """
    Fit ANCOVA (T1 ~ T0 + Treatment) per biomarker and return:
      - stud_flags: rows with |studentized external residual| > 3
      - cook_flags: rows with Cook's D > 4/n
    """
    stud_flags: list[dict] = []
    cook_flags: list[dict] = []

    for bm in BIOMARKERS:
        # Derive T1 column name from delta column (e.g. 'NAD+ Δ' → 'NAD+ T1')
        t1_col = bm.delta_col.replace('Δ', 'T1').strip()
        t0_col = bm.t0_col
        if t1_col not in df.columns or t0_col not in df.columns:
            continue

        subset = df[[t1_col, t0_col, 'Treatment', 'Row', 'Group']].dropna()
        if len(subset) < 10:  # too few observations to fit a reliable model
            continue

        X = sm.add_constant(subset[[t0_col, 'Treatment']])
        model = sm.OLS(subset[t1_col], X).fit()
        influence = OLSInfluence(model)
        stud_res = influence.resid_studentized_external
        cooks_d  = influence.cooks_distance[0]
        threshold = 4 / len(subset)  # Cook's distance threshold: 4/n

        for i, orig_idx in enumerate(subset.index):
            row_id = df.loc[orig_idx, 'Row']
            group  = df.loc[orig_idx, 'Group']
            if abs(stud_res[i]) > 3:
                stud_flags.append({
                    'biomarker': bm.name,
                    'row':       row_id,
                    'group':     group,
                    'residual':  stud_res[i],
                })
            if cooks_d[i] > threshold:
                cook_flags.append({
                    'biomarker': bm.name,
                    'row':       row_id,
                    'group':     group,
                    'cooks_d':   cooks_d[i],
                    'threshold': threshold,
                })

    stud_flags.sort(key=lambda r: abs(r['residual']), reverse=True)
    cook_flags.sort(key=lambda r: r['cooks_d'], reverse=True)
    return stud_flags, cook_flags


# ── commands ──────────────────────────────────────────────────────────────────

def outliers_zscore() -> None:
    """Z-score outliers (|z| > 3) on Δ values → output/a1_outliers_zscore.docx."""
    df = load_data()
    flags = _collect_zscore_flags(df)

    Path('output').mkdir(exist_ok=True)
    doc = _make_doc()

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    _header_row(table, ['Biomarker', 'Row', 'Group', 'Δ value', 'Z-score'])
    for r in flags:
        _data_row(table, [r['biomarker'], str(r['row']), r['group'],
                          f"{r['delta']:.3f}", f"{r['z']:.2f}"])
    set_col_widths(table, [4.0, 1.2, 5.5, 2.2, 2.0])

    doc.add_paragraph()
    _add_legend(doc,
        f'Criterion: |z| > 3 on Δ values across all {len(df)} participants. '
        f'{len(flags)} observation(s) flagged. Outliers retained in the analysis.')

    out = 'output/a1_outliers_zscore.docx'
    doc.save(out)
    print(f'Saved: {out}')


def outliers_studentized() -> None:
    """Studentized residuals (|r*| > 3) from ANCOVA → output/a1_outliers_studentized.docx."""
    df = load_data()
    stud_flags, _ = _collect_ancova_flags(df)

    Path('output').mkdir(exist_ok=True)
    doc = _make_doc()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    _header_row(table, ['Biomarker', 'Row', 'Group', 'Stud. residual'])
    for r in stud_flags:
        _data_row(table, [r['biomarker'], str(r['row']), r['group'],
                          f"{r['residual']:.2f}"])
    set_col_widths(table, [4.0, 1.2, 5.5, 3.0])

    doc.add_paragraph()
    _add_legend(doc,
        f'Criterion: |r*| > 3 (external studentized residual) from ANCOVA model '
        f'T1 ~ T0 + Treatment, fitted per biomarker. '
        f'{len(stud_flags)} observation(s) flagged. Outliers retained in the analysis.')

    out = 'output/a1_outliers_studentized.docx'
    doc.save(out)
    print(f'Saved: {out}')


def outliers_cooks() -> None:
    """Cook's distance (D > 4/n) from ANCOVA → output/a1_outliers_cooks.docx."""
    df = load_data()
    _, cook_flags = _collect_ancova_flags(df)

    Path('output').mkdir(exist_ok=True)
    doc = _make_doc()

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    _header_row(table, ["Biomarker", "Row", "Group", "Cook's D", "Threshold (4/n)"])
    for r in cook_flags:
        _data_row(table, [r['biomarker'], str(r['row']), r['group'],
                          f"{r['cooks_d']:.4f}", f"{r['threshold']:.4f}"])
    set_col_widths(table, [4.0, 1.2, 5.5, 2.2, 2.5])

    doc.add_paragraph()
    _add_legend(doc,
        f"Criterion: Cook's D > 4/n from ANCOVA model T1 ~ T0 + Treatment, "
        f"fitted per biomarker. n varies per biomarker (non-missing observations). "
        f'{len(cook_flags)} observation(s) flagged. Outliers retained in the analysis.')

    out = 'output/a1_outliers_cooks.docx'
    doc.save(out)
    print(f'Saved: {out}')
