from pathlib import Path
from typing import cast

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.styles.style import _ParagraphStyle
from docx.table import _Cell

from biomarkers import AGE_ACCEL_BIOMARKERS, ANALYSIS_BIOMARKERS, LOG_BIOMARKERS
from commands.univariate import _cohens_d, _effect_label, _two_sample_test
from data import load_composites
from docx_utils import (
    C_GRAY, FONT_BODY, FONT_LEGEND, FONT_SUBGROUP,
    add_run, clear_cell, dir_color, new_para, set_col_widths,
)

# Strata are Sex × Age combinations; treatment vs. control is always the comparison axis.
# Each entry: (label, filter_function) where the filter selects a stratum from the full df.
_STRATA: list[tuple[str, object]] = [
    ('all',   lambda df: df),
    ('M',     lambda df: df[df['Sex'] == 'm']),
    ('F',     lambda df: df[df['Sex'] == 'f']),
    ('Y',     lambda df: df[df['Age Group'] == 'Young']),
    ('Avg',   lambda df: df[df['Age Group'] == 'Average']),
    ('O',     lambda df: df[df['Age Group'] == 'Old']),
    ('M×Y',   lambda df: df[(df['Sex'] == 'm') & (df['Age Group'] == 'Young')]),
    ('M×Avg', lambda df: df[(df['Sex'] == 'm') & (df['Age Group'] == 'Average')]),
    ('M×O',   lambda df: df[(df['Sex'] == 'm') & (df['Age Group'] == 'Old')]),
    ('F×Y',   lambda df: df[(df['Sex'] == 'f') & (df['Age Group'] == 'Young')]),
    ('F×Avg', lambda df: df[(df['Sex'] == 'f') & (df['Age Group'] == 'Average')]),
    ('F×O',   lambda df: df[(df['Sex'] == 'f') & (df['Age Group'] == 'Old')]),
]

_GROUP_HEADERS: dict[str, str] = {
    'all': 'Overall / Sex / Age',
    'M×Y': 'Sex × Age',
}


def _test_stratum(
    stratum_df: pd.DataFrame, col: str, min_n: int = 6
) -> tuple[float, float, float]:
    """Run two-sample test (T vs C) within a stratum.

    Returns (diff, p_value, cohens_d); all nan if either group has fewer than min_n observations.
    diff = treatment mean Δ − control mean Δ.
    """
    t_vals = stratum_df[stratum_df['Treatment'] == 1][col].dropna()
    c_vals = stratum_df[stratum_df['Treatment'] == 0][col].dropna()
    if len(t_vals) < min_n or len(c_vals) < min_n:
        return float('nan'), float('nan'), float('nan')
    p, _ = _two_sample_test(t_vals, c_vals)
    d = _cohens_d(t_vals, c_vals)
    diff = float(t_vals.mean() - c_vals.mean())
    return diff, p, d


def _fmt_pd(p: float, d: float) -> str:
    """Format as 'p/d' in 13 characters: p(5)+star(1)+d(6)+trailing_space(1)."""
    if np.isnan(p):
        return '    —        '  # 13 chars
    p_s = '<.001' if p < 0.001 else f'{p:.3f}'
    star = '*' if p < 0.05 else ' '
    d_s = f'{d:+.2f}'
    return f'{p_s}{star}{d_s:>6} '  # 5+1+6+1=13


def univariate_strata() -> None:
    """T vs C Δ comparison within each sex/age stratum for all biomarkers → stdout."""
    df = load_composites()

    _accel_names = {bm.name for bm in AGE_ACCEL_BIOMARKERS}
    _log_names   = {bm.name for bm in LOG_BIOMARKERS}

    def display_name(name: str) -> str:
        if name in _accel_names:
            return f"{name} (res)"
        if name in _log_names:
            return f"{name} (log)"
        return name

    col_w  = 13   # p-value + Cohen's d per stratum
    name_w = 29

    labels = [label for label, _ in _STRATA]

    # Group-header row
    header_line = ' ' * name_w
    group_keys = list(_GROUP_HEADERS.keys())
    for i, label in enumerate(labels):
        if label in _GROUP_HEADERS:
            gh = _GROUP_HEADERS[label]
            next_key = group_keys[group_keys.index(label) + 1] if label != group_keys[-1] else None
            count = (labels.index(next_key) - labels.index(label)) if next_key else (len(labels) - labels.index(label))
            header_line += f'{gh:<{count * col_w}}'
    print(header_line.rstrip())

    # Column-label row: each stratum gets two sub-columns (p and d)
    label_line = f"{'Biomarker':<{name_w}}"
    for label in labels:
        label_line += f'{"p":>6}{"d":>7} '   # 6+6+trailing_space = 13 per stratum
    print(label_line)

    # Sub-header showing stratum names
    sub_line = f"{'':>{name_w}}"
    for label in labels:
        sub_line += f'{label:>13}'
    print(sub_line)
    print('-' * (name_w + col_w * len(labels)))

    for bm in ANALYSIS_BIOMARKERS:
        row = f'{display_name(bm.name):<{name_w}}'
        for _, filt in _STRATA:
            stratum_df = filt(df)
            _, p, d = _test_stratum(stratum_df, bm.delta_col)  # _ = diff (unused in stdout)
            row += _fmt_pd(p, d)
        print(row)

    print()
    print('* p < 0.05  |  — either group n < 6  |  d = Cohen\'s d (pooled SD)')
    print('Strata keys: M=male, F=female, Y=young, Avg=average, O=old')
    print('Each stratum compares treatment Δ vs control Δ using the same test-selection algorithm as univariate.')


# Strata for the docx table: overall + sex + age (no Sex × Age cross-strata).
_STRATA_DOCX: list[tuple[str, object]] = [
    ('Overall', lambda df: df),
    ('M',       lambda df: df[df['Sex'] == 'm']),
    ('F',       lambda df: df[df['Sex'] == 'f']),
    ('Y',       lambda df: df[df['Age Group'] == 'Young']),
    ('Avg',     lambda df: df[df['Age Group'] == 'Average']),
    ('O',       lambda df: df[df['Age Group'] == 'Old']),
]


def _fill_stratum_cell(
    cell: _Cell, diff: float, p: float, d: float, direction: str | None
) -> None:
    """Fill a stratum cell: diff (main, colored when significant), p= d= (secondary gray)."""
    clear_cell(cell)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(0)
    if np.isnan(p):
        add_run(para, '—', FONT_BODY, rgb=C_GRAY)
        return
    sig = p < 0.05
    # Color the diff by direction only when significant; gray otherwise
    val_color = dir_color(diff, direction) if sig else C_GRAY
    add_run(para, f'{diff:+.3f}', FONT_BODY, bold=sig, rgb=val_color)

    p_str = '<.001' if p < 0.001 else f'{p:.3f}'
    p_para = new_para(cell)
    p_para.paragraph_format.space_after = Pt(0)
    add_run(p_para, f'p={p_str}', FONT_SUBGROUP, rgb=C_GRAY)
    d_para = new_para(cell)
    d_para.paragraph_format.space_after = Pt(0)
    add_run(d_para, f'd={d:+.2f}', FONT_SUBGROUP, rgb=C_GRAY)


def univariate_strata_docx() -> None:
    """T vs C Δ comparison within each sex/age stratum for all biomarkers → output/a3_univariate_strata.docx."""
    df = load_composites()
    Path('output').mkdir(exist_ok=True)

    _accel_names = {bm.name for bm in AGE_ACCEL_BIOMARKERS}
    _log_names   = {bm.name for bm in LOG_BIOMARKERS}

    def display_name(name: str) -> str:
        if name in _accel_names:
            return f'{name} (res)'
        if name in _log_names:
            return f'{name} (log)'
        return name

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    # Set default font for the document
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    n_strata = len(_STRATA_DOCX)
    table = doc.add_table(rows=1, cols=1 + n_strata)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for cell in header_cells:
        clear_cell(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

    add_run(header_cells[0].paragraphs[0], 'Biomarker', FONT_BODY, bold=True)
    for i, (label, _) in enumerate(_STRATA_DOCX):
        add_run(header_cells[i + 1].paragraphs[0], label, FONT_BODY, bold=True)

    # Data rows
    for bm in ANALYSIS_BIOMARKERS:
        row_cells = table.add_row().cells
        # Name cell
        clear_cell(row_cells[0])
        name_para = row_cells[0].paragraphs[0]
        name_para.paragraph_format.space_after = Pt(0)
        name_para.paragraph_format.space_before = Pt(0)
        add_run(name_para, display_name(bm.name), FONT_BODY)

        # Stratum cells
        for i, (_, filt) in enumerate(_STRATA_DOCX):
            stratum_df = filt(df)
            diff, p, d = _test_stratum(stratum_df, bm.delta_col)
            _fill_stratum_cell(row_cells[i + 1], diff, p, d, bm.direction)

    # Column widths: name=4.5cm, each stratum=(17-4.5)/6≈2.1cm
    name_w   = 4.5
    strata_w = round((17.0 - name_w) / n_strata, 2)
    set_col_widths(table, [name_w] + [strata_w] * n_strata)

    # Legend
    doc.add_paragraph()
    legend = [
        ('Value',  ': treatment mean Δ − control mean Δ (bold = p < 0.05); — = stratum n < 6 per group.'),
        ('Color',  ': green = favorable change; red = unfavorable. Higher is better: NAD⁺, HDL. Lower is better: all others.'),
        ('p',      ': two-sample test p-value.'),
        ('d',      ': Cohen\'s d (pooled SD).'),
        ('Strata', ': M = male, F = female; Y = Young (<40 yr), Avg = Average (40–54 yr), O = Old (>54 yr).'),
    ]
    for key, val in legend:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, key, FONT_LEGEND, bold=True)
        add_run(para, val, FONT_LEGEND)

    out = 'output/a1_univariate_strata.docx'
    doc.save(out)
    print(f'Saved: {out}')
