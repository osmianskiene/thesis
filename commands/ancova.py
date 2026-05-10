"""ANCOVA: T1 ~ Treatment + T0 [+ Sex] [+ Age] for all biomarkers."""
from dataclasses import dataclass
from pathlib import Path
from typing import cast

import click
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import statsmodels.formula.api as smf
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.styles.style import _ParagraphStyle

from biomarkers import AGE_ACCEL_BIOMARKERS, ANALYSIS_BIOMARKERS, LOG_BIOMARKERS
from data import load_composites
from docx_utils import (
    C_BLACK, C_GRAY, FONT_BODY, FONT_LEGEND,
    add_run, clear_cell, new_para, set_col_widths,
)

# Okabe-Ito colorblind-friendly palette
_TREAT_COLOR   = '#0072B2'  # blue
_CONTROL_COLOR = '#E69F00'  # orange

_ACCEL_NAMES = {bm.name for bm in AGE_ACCEL_BIOMARKERS}
_LOG_NAMES   = {bm.name for bm in LOG_BIOMARKERS}

# Pre-specified endpoint tiers (thesis §5.5.4)
_PRIMARY_NAMES   = frozenset({'Biological Age'})
_SECONDARY_NAMES = frozenset({'NAD⁺', 'Inflammation Score', 'HOMA-IR', 'LDL Cholesterol'})


@dataclass
class _Row:
    """One biomarker's ANCOVA result, collected for optional docx output."""
    name: str
    n: int
    beta: float
    ci_lo: float
    ci_hi: float
    p: float
    r2: float


def _display_name(name: str) -> str:
    if name in _ACCEL_NAMES:
        return f'{name} (res)'
    if name in _LOG_NAMES:
        return f'{name} (log)'
    return name


def _safe_filename(name: str) -> str:
    """Convert biomarker name to a safe filename stem."""
    return name.replace('/', '_').replace(' ', '_').replace('⁺', '+').replace('⁻', '-')


def _tier(bm_name: str) -> str:
    """Return the pre-specified endpoint tier for a biomarker canonical name."""
    if bm_name in _PRIMARY_NAMES:
        return 'Primary'
    if bm_name in _SECONDARY_NAMES:
        return 'Secondary'
    return 'Exploratory'


def _fit_one(
    bm, df: pd.DataFrame, formula: str, sex: bool, age: bool
) -> tuple['_Row', pd.DataFrame, object]:
    """Fit ANCOVA for one biomarker; return (row, fit_df, fitted_model)."""
    bm_t0 = df[bm.t0_col]
    t1    = bm_t0 + df[bm.delta_col]
    # T0 always included in fit_df so the scatter plot can use it as x-axis,
    # even when --t0 is not set (T0 absent from the formula).
    cols: dict = {'T1': t1, 'Treatment': df['Treatment'], 'T0': bm_t0}
    if sex:
        cols['Sex'] = df['Sex_enc']
    if age:
        cols['Age'] = df['Chronological Age T0']
    fit_df = pd.DataFrame(cols).dropna()
    model  = smf.ols(formula, data=fit_df).fit()
    beta   = model.params['Treatment']
    ci     = model.conf_int().loc['Treatment']
    row    = _Row(
        name  = _display_name(bm.name),
        n     = len(fit_df),
        beta  = beta,
        ci_lo = ci[0],
        ci_hi = ci[1],
        p     = model.pvalues['Treatment'],
        r2    = model.rsquared,
    )
    return row, fit_df, model


def _plot_bm(
    bm, row: '_Row', fit_df: pd.DataFrame, model,
    t0: bool, sex: bool, age: bool,
    out_dir: Path, suffix: str,
) -> None:
    """Save Vickers & Altman scatter + regression-line plot for one biomarker."""
    fig, ax = plt.subplots(figsize=(6, 5))

    treat_rows = fit_df[fit_df['Treatment'] == 1]
    ctrl_rows  = fit_df[fit_df['Treatment'] == 0]

    # Individual data points: circles for treatment, plus-signs for control
    ax.scatter(treat_rows['T0'], treat_rows['T1'],
               color=_TREAT_COLOR,   alpha=0.55, s=30, zorder=3,
               marker='o', label='Treatment')
    ax.scatter(ctrl_rows['T0'],  ctrl_rows['T1'],
               color=_CONTROL_COLOR, alpha=0.55, s=30, zorder=3,
               marker='+', label='Control')

    # Square markers at group means (as in Vickers & Altman figure)
    ax.scatter(treat_rows['T0'].mean(), treat_rows['T1'].mean(),
               color=_TREAT_COLOR,   s=100, marker='s', zorder=5, edgecolors='black', linewidths=0.8)
    ax.scatter(ctrl_rows['T0'].mean(),  ctrl_rows['T1'].mean(),
               color=_CONTROL_COLOR, s=100, marker='s', zorder=5, edgecolors='black', linewidths=0.8)

    # Regression lines: vary T0 over observed range, fix optional covariates at sample means
    t0_min, t0_max = fit_df['T0'].min(), fit_df['T0'].max()
    t0_range = np.linspace(t0_min, t0_max, 200)

    p_str = '<.001' if row.p < 0.001 else f'{row.p:.3f}'

    for grp_val, color, linestyle, label in [
        (1, _TREAT_COLOR,   '-',  'Treatment (ANCOVA)'),
        (0, _CONTROL_COLOR, '--', 'Control (ANCOVA)'),
    ]:
        pred_cols: dict = {'Treatment': np.full(200, float(grp_val))}
        if t0:
            pred_cols['T0'] = t0_range
        if sex:
            pred_cols['Sex'] = np.full(200, fit_df['Sex'].mean())
        if age:
            pred_cols['Age'] = np.full(200, fit_df['Age'].mean())
        y_pred = model.predict(pd.DataFrame(pred_cols))
        # When T0 is not a covariate the predicted line is flat; still plotted against t0_range for context
        ax.plot(t0_range, y_pred, color=color, linewidth=2, linestyle=linestyle, label=label)

    ax.set_xlabel(f'{row.name} — pre-treatment (T0)')
    ax.set_ylabel(f'{row.name} — post-treatment (T1)')
    ax.set_title(f'ANCOVA: {row.name}\n'
                 f'β_treat = {row.beta:.3f}  [{row.ci_lo:.3f}, {row.ci_hi:.3f}]  p = {p_str}')

    # Legend: data markers first, then lines
    handles, labels_leg = ax.get_legend_handles_labels()
    order = [0, 1, 2, 3]   # treatment scatter, control scatter, treatment line, control line
    ax.legend([handles[i] for i in order], [labels_leg[i] for i in order],
              fontsize=8, framealpha=0.9)

    fig.tight_layout()
    fname = out_dir / f'ancova{suffix}_{_safe_filename(bm.name)}.png'
    fig.savefig(fname, dpi=150)
    plt.close(fig)


def _save_docx(rows: list['_Row'], formula: str, suffix: str, out_dir: Path) -> None:
    """Write ANCOVA results table to a Word document."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    # Header row
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['Biomarker', 'β treat', 'CI, 95%', 'p', 'R²']
    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, text, FONT_BODY, bold=True)

    # Data rows
    for row in rows:
        p_str = '<.001' if row.p < 0.001 else f'{row.p:.3f}'
        sig   = row.p < 0.05   # bold biomarker name when significant
        cells = table.add_row().cells
        values = [
            row.name,
            f'{row.beta:.3f}',
            f'{row.ci_lo:.3f} – {row.ci_hi:.3f}',
            p_str + (' *' if sig else ''),
            f'{row.r2:.3f}',
        ]
        for cell, text in zip(cells, values):
            clear_cell(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            add_run(para, text, FONT_BODY, bold=(sig and text == row.name))

    #                      name   β     CI    p     R²
    set_col_widths(table, [4.5,  1.8,  3.2,  1.5,  1.5])

    # Legend
    doc.add_paragraph()
    legend = [
        ('Model',   f': {formula} (OLS, Vickers & Altman 2001).'),
        ('β treat', ': adjusted mean difference (treatment − control).'),
        ('CI',      ': 95% confidence interval for the treatment coefficient (OLS regression).'),
        ('R²',      ': proportion of variance in T1 explained by the model.'),
        ('*',       ': p < 0.05 (uncorrected).'),
    ]
    for key, val in legend:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, key, FONT_LEGEND, bold=True)
        add_run(para, val, FONT_LEGEND)

    out = Path('output') / f'ancova{suffix}.docx'
    doc.save(out)
    print(f'Saved: {out}')


def _add_section_row(table, label: str) -> None:
    """Add a shaded full-width section-header row to a table."""
    row    = table.add_row()
    merged = row.cells[0].merge(row.cells[-1])
    clear_cell(merged)
    # Light-gray cell background via OOXML (python-docx has no public shading API)
    tc   = merged._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9D9D9')  # light gray
    tcPr.append(shd)
    para = merged.paragraphs[0]
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, label, FONT_BODY, bold=True)


def _save_docx_primary(
    sections: list[tuple[str, list['_Row']]], formula: str, suffix: str
) -> None:
    """Write ANCOVA results grouped by endpoint tier to a Word document."""
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin  = sec.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['Biomarker', 'β treat', 'CI, 95%', 'p', 'R²']
    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, text, FONT_BODY, bold=True)

    for tier_label, rows in sections:
        _add_section_row(table, tier_label)
        for row in rows:
            p_str = '<.001' if row.p < 0.001 else f'{row.p:.3f}'
            sig   = row.p < 0.05
            cells = table.add_row().cells
            values = [
                row.name,
                f'{row.beta:.3f}',
                f'{row.ci_lo:.3f} – {row.ci_hi:.3f}',
                p_str + (' *' if sig else ''),
                f'{row.r2:.3f}',
            ]
            for cell, text in zip(cells, values):
                clear_cell(cell)
                para = cell.paragraphs[0]
                para.paragraph_format.space_after  = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                add_run(para, text, FONT_BODY, bold=(sig and text == row.name))

    set_col_widths(table, [4.5, 1.8, 3.2, 1.5, 1.5])

    doc.add_paragraph()
    legend = [
        ('Model',   f': {formula} (OLS, Vickers & Altman 2001).'),
        ('β treat', ': adjusted mean difference (treatment − control).'),
        ('CI',      ': 95% confidence interval for the treatment coefficient (OLS regression).'),
        ('R²',      ': proportion of variance in T1 explained by the model.'),
        ('*',       ': p < 0.05 (uncorrected, primary endpoint only).'),
    ]
    for key, val in legend:
        para = doc.add_paragraph()
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        add_run(para, key, FONT_LEGEND, bold=True)
        add_run(para, val, FONT_LEGEND)

    out = Path('output') / f'ancova{suffix}.docx'
    doc.save(out)
    print(f'Saved: {out}')


def ancova(t0: bool = False, sex: bool = False, age: bool = False, docx: bool = False) -> None:
    """ANCOVA for all biomarkers → stdout + output/plots_ancova[_t0][_sex][_age]/."""
    df = load_composites()
    df = df.copy()
    df['Sex_enc'] = (df['Sex'] == 'f').astype(float)   # f=1, m=0

    # Build covariate list and formula from flags
    covariates = []
    if t0:
        covariates.append('T0')
    if sex:
        covariates.append('Sex')
    if age:
        covariates.append('Age')
    formula = 'T1 ~ Treatment' + (' + ' + ' + '.join(covariates) if covariates else '')

    # Output dir and file suffix reflect which covariates are included
    suffix  = ('_t0' if t0 else '') + ('_sex' if sex else '') + ('_age' if age else '')
    out_dir = Path(f'output/plots_ancova{suffix}')
    out_dir.mkdir(parents=True, exist_ok=True)

    header = (f"{'Biomarker':<29} {'n':>4}  "
              f"{'β_treat':>8}  {'95% CI low':>10}  {'95% CI high':>11}  "
              f"{'p':>7}  {'R²':>6}")
    print(header)
    print('-' * len(header))

    rows: list[_Row] = []   # collected for optional docx output

    for bm in ANALYSIS_BIOMARKERS:
        row, fit_df, model = _fit_one(bm, df, formula, sex, age)
        rows.append(row)
        _plot_bm(bm, row, fit_df, model, t0, sex, age, out_dir, suffix)

        p_str = '<.001' if row.p < 0.001 else f'{row.p:.3f}'
        sig   = '*' if row.p < 0.05 else ' '
        print(f'{row.name:<29} {row.n:>4}  {row.beta:>8.3f}  {row.ci_lo:>10.3f}  {row.ci_hi:>11.3f}  '
              f'{p_str:>6}{sig}  {row.r2:>6.3f}')

    print()
    print('* p < 0.05')
    print(f'Model: {formula} (OLS, Vickers & Altman 2001)')
    print(f'β_treat = adjusted mean difference (treatment − control), 95% CI, p-value from OLS t-test')
    print(f'Plots saved to {out_dir}/')

    if docx:
        _save_docx(rows, formula, suffix, out_dir)


def ancova_primary(t0: bool = False, sex: bool = False, age: bool = False, docx: bool = False) -> None:
    """ANCOVA grouped by endpoint tier (primary / secondary / exploratory) → stdout + plots."""
    df = load_composites()
    df = df.copy()
    df['Sex_enc'] = (df['Sex'] == 'f').astype(float)   # f=1, m=0

    covariates = []
    if t0:
        covariates.append('T0')
    if sex:
        covariates.append('Sex')
    if age:
        covariates.append('Age')
    formula = 'T1 ~ Treatment' + (' + ' + ' + '.join(covariates) if covariates else '')

    suffix  = '_primary' + ('_t0' if t0 else '') + ('_sex' if sex else '') + ('_age' if age else '')
    out_dir = Path(f'output/plots_ancova{suffix}')
    out_dir.mkdir(parents=True, exist_ok=True)

    tier_rows: dict[str, list[_Row]] = {'Primary': [], 'Secondary': [], 'Exploratory': []}

    for bm in ANALYSIS_BIOMARKERS:
        row, fit_df, model = _fit_one(bm, df, formula, sex, age)
        tier_rows[_tier(bm.name)].append(row)
        _plot_bm(bm, row, fit_df, model, t0, sex, age, out_dir, suffix)

    header = (f"{'Biomarker':<29} {'n':>4}  "
              f"{'β_treat':>8}  {'95% CI low':>10}  {'95% CI high':>11}  "
              f"{'p':>7}  {'R²':>6}")

    for tier_label in ('Primary', 'Secondary', 'Exploratory'):
        print(f'\n{tier_label.upper()}')
        print(header)
        print('-' * len(header))
        for row in tier_rows[tier_label]:
            p_str = '<.001' if row.p < 0.001 else f'{row.p:.3f}'
            sig   = '*' if row.p < 0.05 else ' '
            print(f'{row.name:<29} {row.n:>4}  {row.beta:>8.3f}  {row.ci_lo:>10.3f}  {row.ci_hi:>11.3f}  '
                  f'{p_str:>6}{sig}  {row.r2:>6.3f}')

    print()
    print('* p < 0.05')
    print(f'Model: {formula} (OLS, Vickers & Altman 2001)')
    print(f'β_treat = adjusted mean difference (treatment − control), 95% CI, p-value from OLS t-test')
    print(f'Plots saved to {out_dir}/')

    if docx:
        sections = [(tier, tier_rows[tier]) for tier in ('Primary', 'Secondary', 'Exploratory')]
        _save_docx_primary(sections, formula, suffix)


@click.command()
@click.option('--t0',      is_flag=True, help='Include T0 (baseline) as covariate')
@click.option('--sex',     is_flag=True, help='Include Sex as covariate')
@click.option('--age',     is_flag=True, help='Include Age as covariate')
@click.option('--docx',    is_flag=True, help='Also save results table as output/ancova[...].docx')
@click.option('--primary', is_flag=True, help='Group results by endpoint tier (primary/secondary/exploratory)')
def main(t0: bool, sex: bool, age: bool, docx: bool, primary: bool) -> None:
    """ANCOVA (T1 ~ Treatment [+ T0] [+ Sex] [+ Age]) for all biomarkers."""
    if primary:
        ancova_primary(t0=t0, sex=sex, age=age, docx=docx)
    else:
        ancova(t0=t0, sex=sex, age=age, docx=docx)
