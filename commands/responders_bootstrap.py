from pathlib import Path
from typing import cast

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from matplotlib.lines import Line2D

from commands.responders import DOMAIN_BIOMARKERS, DOMAINS
from data import load_composites
from docx_utils import C_GRAY, C_GREEN, C_RED, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

_OUT_DIR = Path('output/plots_bootstrap')
_N_BOOT  = 2_000


class BootstrapResult:
    """Per-biomarker bootstrap CI result."""
    __slots__ = ('name', 'direction', 'domain', 'obs_diff', 'ci_lo', 'ci_hi',
                 'excludes_zero', 'favorable')

    def __init__(self, name: str, direction: str, domain: str,
                 obs_diff: float, ci_lo: float, ci_hi: float) -> None:
        self.name         = name
        self.direction    = direction
        self.domain       = domain
        self.obs_diff     = obs_diff
        self.ci_lo        = ci_lo
        self.ci_hi        = ci_hi
        self.excludes_zero = (ci_lo > 0) or (ci_hi < 0)
        # Favorable: Treatment moved in the beneficial direction relative to Control
        self.favorable    = (
            (direction == 'down' and obs_diff < 0) or
            (direction == 'up'   and obs_diff > 0)
        )


def responders_bootstrap() -> None:
    """Bootstrap 95% CIs for mean T−C Δ per biomarker → stdout + output/a3_bootstrap_ci.docx + forest plot."""
    df  = load_composites()
    rng = np.random.default_rng(42)

    # Build name → domain lookup from DOMAINS
    bm_to_domain: dict[str, str] = {}
    for domain in DOMAINS:
        for bm_name in domain.biomarker_names:
            bm_to_domain[bm_name] = domain.name

    results: list[BootstrapResult] = []
    for bm in DOMAIN_BIOMARKERS:
        t_vals = df.loc[df['Treatment'] == 1, bm.delta_col].dropna().to_numpy()
        c_vals = df.loc[df['Treatment'] == 0, bm.delta_col].dropna().to_numpy()
        obs_diff = float(t_vals.mean() - c_vals.mean())

        boot_diffs = np.empty(_N_BOOT)
        for k in range(_N_BOOT):
            t_boot = rng.choice(t_vals, size=len(t_vals), replace=True)
            c_boot = rng.choice(c_vals, size=len(c_vals), replace=True)
            boot_diffs[k] = t_boot.mean() - c_boot.mean()

        ci_lo = float(np.percentile(boot_diffs, 2.5))
        ci_hi = float(np.percentile(boot_diffs, 97.5))
        domain = bm_to_domain.get(bm.name, '')
        results.append(BootstrapResult(bm.name, bm.direction or '', domain,
                                       obs_diff, ci_lo, ci_hi))

    # ── CLI output ─────────────────────────────────────────────────────────────
    print(f'Bootstrap 95% CIs – Mean Δ Difference (Treatment − Control)')
    print(f'  {_N_BOOT:,} resamples, within-group sampling with replacement')
    print('=' * 75)
    print(f"{'Biomarker':<22} {'Dir':>4} {'Obs diff':>10} {'95% CI':>22}  Reliable  Favorable")
    print('-' * 75)

    for r in results:
        rel = 'YES' if r.excludes_zero else ''
        fav = 'YES' if r.favorable    else ''
        print(f'{r.name:<22} {r.direction:>4} {r.obs_diff:>10.4f} '
              f'[{r.ci_lo:>8.4f}, {r.ci_hi:>8.4f}]  {rel:<8}  {fav}')

    n_reliable          = sum(1 for r in results if r.excludes_zero)
    n_favorable_reliable = sum(1 for r in results if r.excludes_zero and r.favorable)
    print(f'\n  {n_reliable}/{len(results)} biomarkers with CI excluding zero')
    print(f'  {n_favorable_reliable}/{n_reliable} of those in the favorable direction')

    _write_docx(results)
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    _plot_forest(results)
    print(f'\nPlot saved to {_OUT_DIR}/')


def _write_docx(results: list[BootstrapResult]) -> None:
    """Save bootstrap CI table → output/a3_bootstrap_ci.docx."""
    out = Path('output/a3_bootstrap_ci.docx')
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['Biomarker', 'Domain', 'Dir.', 'Obs. diff.', 'CI lower', 'CI upper', 'Sig.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_col_widths(table, [3.5, 3.0, 1.2, 2.2, 2.2, 2.2, 1.2])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    for r in results:
        sig = '*' if r.excludes_zero else ''
        # Green * = CI excludes zero and favorable; red * = excludes zero but unfavorable
        sig_color = (C_GREEN if r.favorable else C_RED) if r.excludes_zero else None
        row = table.add_row()
        for cell in row.cells:
            clear_cell(cell)
        vals_plain = [r.name, r.domain, r.direction,
                      f'{r.obs_diff:.4f}', f'{r.ci_lo:.4f}', f'{r.ci_hi:.4f}']
        for cell, val in zip(row.cells, vals_plain):
            add_run(cell.paragraphs[0], val, FONT_BODY)
        if sig_color is not None:
            add_run(row.cells[-1].paragraphs[0], sig, FONT_BODY, rgb=sig_color)
        else:
            add_run(row.cells[-1].paragraphs[0], sig, FONT_BODY)

    legend = (
        f'Bootstrap 95% confidence intervals for mean Δ difference (Treatment − Control) '
        f'per biomarker ({_N_BOOT:,} resamples, within-group sampling with replacement). '
        f'Dir. = expected improvement direction. '
        f'* CI excludes zero (reliable directional effect). '
        f'Green * = reliable and in the favorable direction; '
        f'Red * = reliable but in the unfavorable direction.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(str(out))
    print(f'Saved {out}')


def _plot_forest(results: list[BootstrapResult]) -> None:
    """Forest plot: observed mean T−C difference with 95% bootstrap CI per biomarker.

    Biomarkers ordered by domain, coloured by reliability and direction:
      green  = CI excludes zero and favorable
      orange = CI excludes zero but unfavorable
      gray   = CI includes zero
    """
    n   = len(results)
    fig, ax = plt.subplots(figsize=(8, max(5, n * 0.42)))

    # Build domain separator positions (y-axis tick labels include domain header rows)
    y_labels: list[str]  = []
    y_positions: list[int] = []
    plot_items: list[BootstrapResult | None] = []  # None = domain header row

    prev_domain = ''
    y = 0
    for r in results:
        if r.domain != prev_domain:
            y_labels.append(f'— {r.domain} —')
            y_positions.append(y)
            plot_items.append(None)
            y += 1
            prev_domain = r.domain
        y_labels.append(r.name)
        y_positions.append(y)
        plot_items.append(r)
        y += 1

    total_rows = y
    ax.set_ylim(-0.5, total_rows - 0.5)

    for yp, item in zip(y_positions, plot_items):
        if item is None:
            continue
        if item.excludes_zero and item.favorable:
            color = '#009E73'    # Okabe-Ito green
        elif item.excludes_zero:
            color = '#E69F00'    # Okabe-Ito orange (unfavorable but reliable)
        else:
            color = '#999999'   # gray
        ax.plot([item.ci_lo, item.ci_hi], [yp, yp], color=color, linewidth=1.8)
        ax.scatter([item.obs_diff], [yp], color=color, s=45, zorder=3)

    ax.axvline(0, color='black', linewidth=0.8, linestyle='--')
    ax.set_yticks(y_positions)
    ax.set_yticklabels(y_labels, fontsize=8)
    ax.set_xlabel('Mean Δ difference (Treatment − Control)', fontsize=10)
    ax.set_title(f'Bootstrap 95% CIs – Mean Δ Difference\n({_N_BOOT:,} resamples per biomarker)',
                 fontsize=10, fontweight='bold')

    legend_handles = [
        Line2D([0], [0], color='#009E73', linewidth=2, marker='o', markersize=6,
               label='CI excludes zero, favorable'),
        Line2D([0], [0], color='#E69F00', linewidth=2, marker='o', markersize=6,
               label='CI excludes zero, unfavorable'),
        Line2D([0], [0], color='#999999', linewidth=2, marker='o', markersize=6,
               label='CI includes zero'),
    ]
    ax.legend(handles=legend_handles, fontsize=8, loc='lower right')

    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'bootstrap_ci_forest.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {_OUT_DIR}/bootstrap_ci_forest.png')
