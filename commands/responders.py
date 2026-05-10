from functools import cache
from pathlib import Path
from typing import cast, NamedTuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from scipy.stats import fisher_exact, mannwhitneyu

from biomarkers import ANALYSIS_BIOMARKERS, Biomarker
from data import load_composites
from docx_utils import C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

# Okabe-Ito: blue = Treatment, orange = Control
_COLOR_TREAT = '#0072B2'
_COLOR_CTRL  = '#E69F00'

_OUT_DIR = Path('output/plots_responders')


class Domain(NamedTuple):
    name: str
    biomarker_names: list[str]  # display names matching ANALYSIS_BIOMARKERS entries


# Five clinically coherent biomarker domains reflecting TLM01's proposed mechanism
DOMAINS: list[Domain] = [
    Domain('Inflammation',      ['CRP', 'IL-6', 'Fibrinogen', 'Inflammation Score']),
    Domain('Glycemic health',   ['Glucose', 'Insulin', 'HOMA-IR', 'HbA1c']),
    Domain('Lipid profile',     ['Triglycerides', 'HDL Cholesterol', 'LDL Cholesterol', 'Total Cholesterol']),
    Domain('Biological ageing', ['Biological Age', 'Eye Age', 'Hearing Age', 'Memory Age']),
    Domain('Cellular energy',   ['NAD⁺']),
]

# AFP, LDH, and Non-HDL excluded: non-specific markers or redundant with other lipids
_EXCLUDED_NAMES = {'AFP', 'LDH', 'Non-HDL Cholesterol'}

# 17 domain-relevant biomarkers used for both ML classification and bootstrap CIs
DOMAIN_BIOMARKERS: list[Biomarker] = [
    bm for bm in ANALYSIS_BIOMARKERS if bm.name not in _EXCLUDED_NAMES
]


@cache
def compute_responders() -> pd.DataFrame:
    """Compute per-domain and overall responder status for all participants.

    Returns a DataFrame (same index as patient_data) with columns:
      domain_{Domain.name} (bool): True if ≥1 biomarker in domain improved
      domain_count (int): number of domains met (0–5)
      strict_responder (bool): all 5/5 domains met
      lenient_responder (bool): ≥3/5 domains met
      response_score (int): count of improved individual markers out of 17
      Treatment (int): group label carried from load_composites
    """
    df = load_composites()
    bm_lookup = {bm.name: bm for bm in DOMAIN_BIOMARKERS}
    result = pd.DataFrame(index=df.index)

    for domain in DOMAINS:
        any_improved = pd.Series(False, index=df.index)
        for bm_name in domain.biomarker_names:
            bm = bm_lookup[bm_name]
            delta = df[bm.delta_col]
            improved = (delta < 0) if bm.direction == 'down' else (delta > 0)
            # fillna(False): missing value does not count as improvement
            any_improved = any_improved | improved.fillna(False)
        result[f'domain_{domain.name}'] = any_improved

    domain_cols = [f'domain_{d.name}' for d in DOMAINS]
    result['domain_count'] = result[domain_cols].sum(axis=1)
    result['strict_responder'] = result['domain_count'] == 5
    result['lenient_responder'] = result['domain_count'] >= 3

    # Count total improved markers across all 17 biomarkers
    score = pd.Series(0, index=df.index, dtype=int)
    for bm in DOMAIN_BIOMARKERS:
        delta = df[bm.delta_col]
        improved = (delta < 0) if bm.direction == 'down' else (delta > 0)
        score += improved.fillna(False).astype(int)
    result['response_score'] = score

    result['Treatment'] = df['Treatment']
    return result


def _fisher(resp: pd.Series, treatment: pd.Series) -> tuple[float, float]:
    """Fisher's exact test (two-sided): responder rate Treatment vs Control."""
    a = int(((treatment == 1) &  resp).sum())   # T responders
    b = int(((treatment == 1) & ~resp).sum())   # T non-responders
    c = int(((treatment == 0) &  resp).sum())   # C responders
    d = int(((treatment == 0) & ~resp).sum())   # C non-responders
    odds_ratio, p = fisher_exact([[a, b], [c, d]], alternative='two-sided')
    return float(odds_ratio), float(p)


def responders_classify() -> None:
    """Multi-domain responder classification → stdout + output/a3_responders.docx + plots."""
    resp = compute_responders()
    treat = resp['Treatment']
    n_t = int((treat == 1).sum())
    n_c = int((treat == 0).sum())

    print('Multi-Domain Responder Analysis')
    print('=' * 65)
    print(f'  Treatment n = {n_t},  Control n = {n_c}')
    print(f'  Strict responder  = all 5/5 domains met')
    print(f'  Lenient responder = ≥3/5 domains met')
    print()
    print(f"{'Domain':<22} {'T rate':>8} {'C rate':>8} {'OR':>7} {'p':>8}  Sig.")
    print('-' * 65)

    domain_rows: list[tuple] = []
    for domain in DOMAINS:
        col = f'domain_{domain.name}'
        t_rate = float(resp.loc[treat == 1, col].mean())
        c_rate = float(resp.loc[treat == 0, col].mean())
        or_, p = _fisher(resp[col], treat)
        sig = '*' if p < 0.05 else ''
        domain_rows.append((domain.name, t_rate, c_rate, or_, p, sig))
        print(f'{domain.name:<22} {t_rate:>7.1%} {c_rate:>8.1%} {or_:>7.2f} {p:>8.4f}  {sig}')

    print()
    overall_rows: list[tuple] = []
    for label, col in [('Strict (5/5)', 'strict_responder'), ('Lenient (≥3/5)', 'lenient_responder')]:
        t_rate = float(resp.loc[treat == 1, col].mean())
        c_rate = float(resp.loc[treat == 0, col].mean())
        or_, p = _fisher(resp[col], treat)
        sig = '*' if p < 0.05 else ''
        overall_rows.append((label, t_rate, c_rate, or_, p, sig))
        print(f'{label:<22} {t_rate:>7.1%} {c_rate:>8.1%} {or_:>7.2f} {p:>8.4f}  {sig}')

    t_score = resp.loc[treat == 1, 'response_score']
    c_score = resp.loc[treat == 0, 'response_score']
    u_stat, p_score = mannwhitneyu(t_score, c_score, alternative='two-sided')
    print()
    print('Response score (0–17 improved markers):')
    print(f'  Treatment: {t_score.mean():.2f} ± {t_score.std():.2f}')
    print(f'  Control:   {c_score.mean():.2f} ± {c_score.std():.2f}')
    print(f'  Mann-Whitney U = {u_stat:.1f}, p = {p_score:.4f}  {"*" if p_score < 0.05 else ""}')
    print('* p < 0.05')

    _write_docx(domain_rows, overall_rows, u_stat, p_score, n_t, n_c)
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    _plot_responder_rates(domain_rows, overall_rows)
    _plot_domain_count(resp, treat, n_t, n_c)
    print(f'\nPlots saved to {_OUT_DIR}/')


def _write_docx(
    domain_rows: list[tuple],
    overall_rows: list[tuple],
    score_u: float,
    score_p: float,
    n_t: int,
    n_c: int,
) -> None:
    """Save responder classification table → output/a3_responders.docx."""
    out = Path('output/a3_responders.docx')
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['Criterion', 'T rate', 'C rate', 'Odds ratio', 'p', 'Sig.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_col_widths(table, [4.5, 2.0, 2.0, 2.5, 2.0, 1.2])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    def _add_data_row(label: str, t_rate: float, c_rate: float,
                      or_: float, p: float, sig: str) -> None:
        row = table.add_row()
        for cell in row.cells:
            clear_cell(cell)
        for cell, val in zip(row.cells,
                              [label, f'{t_rate:.1%}', f'{c_rate:.1%}',
                               f'{or_:.2f}', f'{p:.4f}', sig]):
            add_run(cell.paragraphs[0], val, FONT_BODY)

    for name, t_rate, c_rate, or_, p, sig in domain_rows:
        _add_data_row(name, t_rate, c_rate, or_, p, sig)

    sep = table.add_row()
    for cell in sep.cells:
        clear_cell(cell)
    add_run(sep.cells[0].paragraphs[0], 'Overall responder status', FONT_BODY, bold=True)

    for label, t_rate, c_rate, or_, p, sig in overall_rows:
        _add_data_row(label, t_rate, c_rate, or_, p, sig)

    legend = (
        f'Multi-domain responder rates (Treatment n={n_t}, Control n={n_c}). '
        f'Domain-level responder: improvement in ≥1 biomarker within that domain '
        f'in the clinically expected direction. '
        f'Strict responder: all 5/5 domains met. Lenient responder: ≥3/5 domains. '
        f"Odds ratio from Fisher's exact test (two-sided). "
        f'Response score Mann-Whitney U = {score_u:.1f}, p = {score_p:.4f}. '
        f'* p < 0.05.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(str(out))
    print(f'Saved {out}')


def _plot_responder_rates(
    domain_rows: list[tuple],
    overall_rows: list[tuple],
) -> None:
    """Left: overall strict/lenient rates; right: domain-level rates, grouped bars T vs C."""
    fig, (ax_main, ax_dom) = plt.subplots(1, 2, figsize=(13, 5))
    w = 0.35

    # ── Overall (left) ─────────────────────────────────────────────────────────
    x = np.arange(len(overall_rows))
    ax_main.bar(x - w/2, [r[1] * 100 for r in overall_rows], w,
                color=_COLOR_TREAT, label='Treatment', alpha=0.85)
    ax_main.bar(x + w/2, [r[2] * 100 for r in overall_rows], w,
                color=_COLOR_CTRL, label='Control', alpha=0.85)
    for i, row in enumerate(overall_rows):
        p = row[4]
        ann = '* p<0.05' if p < 0.05 else f'p={p:.2f}'
        ax_main.text(x[i], max(row[1], row[2]) * 100 + 3, ann,
                     ha='center', va='bottom', fontsize=8)
    ax_main.set_xticks(x)
    ax_main.set_xticklabels([r[0] for r in overall_rows], fontsize=9)
    ax_main.set_ylabel('Responder rate (%)', fontsize=10)
    ax_main.set_title('Overall Responder Rates', fontsize=10, fontweight='bold')
    ax_main.set_ylim(0, 110)
    ax_main.legend(fontsize=8)

    # ── Domain breakdown (right, horizontal bars) ──────────────────────────────
    y = np.arange(len(domain_rows))
    ax_dom.barh(y - w/2, [r[1] * 100 for r in domain_rows], w,
                color=_COLOR_TREAT, label='Treatment', alpha=0.85)
    ax_dom.barh(y + w/2, [r[2] * 100 for r in domain_rows], w,
                color=_COLOR_CTRL, label='Control', alpha=0.85)
    for i, row in enumerate(domain_rows):
        if row[4] < 0.05:
            ax_dom.text(max(row[1], row[2]) * 100 + 1.5, y[i], '*',
                        va='center', fontsize=11)
    ax_dom.set_yticks(y)
    ax_dom.set_yticklabels([r[0] for r in domain_rows], fontsize=9)
    ax_dom.set_xlabel('Responder rate (%)', fontsize=10)
    ax_dom.set_title('Domain-Level Responder Rates', fontsize=10, fontweight='bold')
    ax_dom.set_xlim(0, 115)
    ax_dom.legend(fontsize=8)

    fig.suptitle('Multi-Domain Responder Analysis – Treatment vs Control', fontsize=11, y=1.01)
    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'responder_rates.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {_OUT_DIR}/responder_rates.png')


def _plot_domain_count(
    resp: pd.DataFrame,
    treat: pd.Series,
    n_t: int,
    n_c: int,
) -> None:
    """Stacked bar: distribution of number of domains met (0–5) per group."""
    # Okabe-Ito extended for 6 levels
    colors = ['#E69F00', '#56B4E9', '#009E73', '#F0E442', '#0072B2', '#D55E00']
    groups = {'Treatment': treat == 1, 'Control': treat == 0}
    totals = {'Treatment': n_t, 'Control': n_c}

    fig, ax = plt.subplots(figsize=(6, 5))
    x = np.arange(len(groups))
    bottoms = np.zeros(len(groups))

    for count in range(6):
        heights = [
            (resp.loc[mask, 'domain_count'] == count).sum() / totals[name] * 100
            for name, mask in groups.items()
        ]
        ax.bar(x, heights, bottom=bottoms, color=colors[count],
               label=f'{count}/5 domains', alpha=0.9)
        bottoms += np.array(heights)

    ax.set_xticks(x)
    ax.set_xticklabels(list(groups.keys()), fontsize=10)
    ax.set_ylabel('Participants (%)', fontsize=10)
    ax.set_title('Distribution of Domain Count Met\nper Group', fontsize=10, fontweight='bold')
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc='upper left')
    ax.set_ylim(0, 105)

    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'domain_count_distribution.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {_OUT_DIR}/domain_count_distribution.png')
