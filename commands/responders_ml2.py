from pathlib import Path
from typing import NamedTuple, cast

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

from commands.responders import DOMAIN_BIOMARKERS, compute_responders
from data import load_composites
from docx_utils import C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

# Okabe-Ito colors for three classifiers
_CLF_COLORS = {
    'Random Forest':       '#0072B2',
    'Gradient Boosting':   '#009E73',
    'Logistic Regression': '#D55E00',
}

_N_PERM = 200

# Biological Age is always excluded from predictors (it is the default outcome)
_EXCLUDED_PREDICTOR = 'Biological Age'

# Non-biomarker predictors that are valid --only values
_EXTRA_PREDICTORS = ('Sex', 'Chronological Age')

# All valid --only names: biomarker names (excluding Biological Age) + extra predictors
VALID_ONLY_VALUES: list[str] = (
    [bm.name for bm in DOMAIN_BIOMARKERS if bm.name != _EXCLUDED_PREDICTOR]
    + list(_EXTRA_PREDICTORS)
)


class Target(NamedTuple):
    name: str              # display name used in output titles and legends
    slug: str              # appended to output filenames; empty string = default target
    responder_col: str | None  # column in compute_responders(); None = use Biological Age AgeAccel Δ < 0


# Default and all named targets
TARGET_BIOLOGICAL_AGE = Target('Biological Age improvement',        'biological_age',             None)
TARGET_STRICT         = Target('Strict responder (5/5 domains)',    'strict_responder',           'strict_responder')
TARGET_INFLAMMATION   = Target('Inflammation domain responder',     'inflammation_responder',     'domain_Inflammation')
TARGET_GLYCEMIC       = Target('Glycemic health domain responder',  'glycemic_health_responder',  'domain_Glycemic health')
TARGET_LIPID          = Target('Lipid profile domain responder',    'lipid_profile_responder',    'domain_Lipid profile')
TARGET_AGEING         = Target('Biological ageing domain responder','biological_ageing_responder','domain_Biological ageing')
TARGET_CELLULAR       = Target('Cellular energy domain responder',  'cellular_energy_responder',  'domain_Cellular energy')


class ClassifierResult(NamedTuple):
    name: str       # classifier display name
    mean_auc: float
    sd_auc: float
    perm_p: float
    sig: str        # '*' if p < 0.05, else ''


def _slug(names: list[str]) -> str:
    """'Biological Age', 'NAD+' → 'biological_age_nad'"""
    import re
    return '_'.join(re.sub(r'[^a-z0-9]+', '_', n.lower()).strip('_') for n in names)


def _out_paths(only: tuple[str, ...], selected_names: list[str], target: Target) -> tuple[Path, Path]:
    """Return (plot_dir, docx_path) with predictor and target slug suffixes."""
    pred_suffix = f'_{_slug(selected_names)}' if only else ''
    tgt_suffix  = f'_{target.slug}'
    suffix    = pred_suffix + tgt_suffix
    plot_dir  = Path(f'output/plots_responders_ml2{suffix}')
    docx_path = Path(f'output/a3_ml2{suffix}.docx')
    return plot_dir, docx_path


def _compute_y(df: pd.DataFrame, target: Target) -> pd.Series:
    """Return a binary (0/1) outcome Series aligned to df.index."""
    if target.responder_col is None:
        # Biological Age AgeAccel Δ < 0 means age acceleration decreased (improved)
        return (df['Biological Age AgeAccel Δ'] < 0).astype(int)
    resp = compute_responders()
    return resp[target.responder_col].reindex(df.index).astype(int)


def responders_ml2(
    only: tuple[str, ...] = (),
    target: Target = TARGET_BIOLOGICAL_AGE,
    test: bool = False,
) -> None:
    """ML prediction of a binary outcome from T0 biomarkers + optional sex/age → stdout + plots + docx."""
    df = load_composites()

    # Validate --only values against all known predictor names
    if only:
        unknown = set(only) - set(VALID_ONLY_VALUES)
        if unknown:
            raise ValueError(
                f'Unknown --only value(s): {sorted(unknown)}\n'
                f'Valid values: {VALID_ONLY_VALUES}'
            )

    # Biomarker predictors: all DOMAIN_BIOMARKERS except the excluded one; filtered by --only when active
    predictor_biomarkers = [
        bm for bm in DOMAIN_BIOMARKERS
        if bm.name != _EXCLUDED_PREDICTOR and (not only or bm.name in only)
    ]
    # Sex and Chronological Age: always included unless --only is active and omits them
    include_sex = not only or 'Sex' in only
    include_age = not only or 'Chronological Age' in only

    if not predictor_biomarkers and not include_sex and not include_age:
        raise ValueError('No predictors selected; check --only values.')

    t0_cols  = [bm.t0_col for bm in predictor_biomarkers]
    bm_names = [bm.name   for bm in predictor_biomarkers]

    # output paths encode both the predictor filter and the chosen target
    selected_names = bm_names + (['Sex'] if include_sex else []) + (['Chronological Age'] if include_age else [])
    out_dir, docx_path = _out_paths(only, selected_names, target)

    # Encode sex from the full dataset so treatment and control share the same binary mapping
    sex_dummies_all = pd.get_dummies(df['Sex'], drop_first=True)
    sex_encoded_all = sex_dummies_all.iloc[:, 0].astype(float)  # single binary column

    df_ctrl  = df[df['Treatment'] == 0]
    df_treat = df[df['Treatment'] == 1]

    feature_names = bm_names + (['Sex'] if include_sex else []) + (['Chronological Age'] if include_age else [])

    def _build_features(subset: pd.DataFrame) -> pd.DataFrame:
        """Assemble the feature DataFrame for a given participant subset."""
        feat = subset[t0_cols].copy()
        if include_sex:
            feat['Sex'] = sex_encoded_all.reindex(subset.index)
        if include_age:
            feat['Chronological Age'] = subset['Chronological Age T0']
        return feat

    if test:
        # Hold out the non-training treatment patients for evaluation instead of control group
        df_train    = df_treat.sample(n=70, random_state=42)
        df_eval     = df_treat.drop(df_train.index)
        eval_label  = 'Hold-out treatment'
    else:
        df_train    = df_treat
        df_eval     = df_ctrl
        eval_label  = 'Control'

    y_raw      = _compute_y(df_train, target)
    feature_df = _build_features(df_train)

    # Drop rows with any missing feature or missing outcome
    valid = feature_df.notna().all(axis=1) & y_raw.notna()
    X = feature_df[valid].to_numpy()
    y = y_raw[valid].to_numpy()

    n_pos   = int(y.sum())
    n_total = len(y)

    print(f'ML Prediction: {target.name} ({len(feature_names)} predictors)')
    print('=' * 65)
    cohort_note = f'treatment group, 70-patient train split' if test else f'treatment group only'
    print(f'  Cohort: {cohort_note} (n={n_total})')
    if only:
        print(f'  Filter: {", ".join(bm_names)}')
    print(f'  n = {n_total}  (Positive: {n_pos}, Negative: {n_total - n_pos})')
    extras = ', '.join(x for x, inc in [('Sex', include_sex), ('Chronological Age', include_age)] if inc)
    print(f'  Predictors: {len(bm_names)} biomarker T0 value(s)' + (f' + {extras}' if extras else ''))
    print(f'  Cross-validation: stratified 5-fold, metric: AUC-ROC')
    print(f'  Permutation test: {_N_PERM} permutations (shuffled labels, same CV)')
    print()
    print(f"{'Classifier':<22} {'Mean AUC':>10} {'SD':>6}  Perm p  Sig.")
    print('-' * 55)

    classifiers: dict[str, Pipeline] = {
        'Random Forest': Pipeline([
            ('scaler', StandardScaler()),
            ('clf', RandomForestClassifier(n_estimators=500, random_state=42,
                                           class_weight='balanced')),
        ]),
        'Gradient Boosting': Pipeline([
            ('scaler', StandardScaler()),
            # GradientBoostingClassifier has no class_weight; imbalance noted in thesis
            ('clf', GradientBoostingClassifier(n_estimators=200, random_state=42)),
        ]),
        'Logistic Regression': Pipeline([
            ('scaler', StandardScaler()),
            ('clf', LogisticRegression(max_iter=1000, random_state=42,
                                       class_weight='balanced')),
        ]),
    }

    cv  = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
    rng = np.random.default_rng(42)

    cv_aucs:  dict[str, list[float]]      = {}
    all_fpr:  dict[str, list[np.ndarray]] = {}
    all_tpr:  dict[str, list[np.ndarray]] = {}
    rf_importances: np.ndarray | None     = None
    clf_results: list[ClassifierResult]   = []
    best_pipe: Pipeline | None            = None
    best_clf_name: str                    = ''
    best_mean_auc: float                  = -1.0

    for clf_name, pipe in classifiers.items():
        fold_aucs: list[float]      = []
        fold_fpr:  list[np.ndarray] = []
        fold_tpr:  list[np.ndarray] = []

        for train_idx, test_idx in cv.split(X, y):
            pipe.fit(X[train_idx], y[train_idx])
            y_prob = pipe.predict_proba(X[test_idx])[:, 1]
            fold_aucs.append(float(roc_auc_score(y[test_idx], y_prob)))
            fpr, tpr, _ = roc_curve(y[test_idx], y_prob)
            fold_fpr.append(fpr)
            fold_tpr.append(tpr)

        cv_aucs[clf_name] = fold_aucs
        all_fpr[clf_name] = fold_fpr
        all_tpr[clf_name] = fold_tpr
        mean_obs = float(np.mean(fold_aucs))
        sd_obs   = float(np.std(fold_aucs))

        # Permutation test: shuffle labels, re-run CV, compare mean AUC
        null_aucs: list[float] = []
        for _ in range(_N_PERM):
            y_perm = rng.permutation(y)
            perm_fold: list[float] = []
            for train_idx, test_idx in cv.split(X, y_perm):
                pipe.fit(X[train_idx], y_perm[train_idx])
                y_prob_p = pipe.predict_proba(X[test_idx])[:, 1]
                perm_fold.append(float(roc_auc_score(y_perm[test_idx], y_prob_p)))
            null_aucs.append(float(np.mean(perm_fold)))

        # +1/(N+1) continuity correction avoids p = 0
        perm_p = float((np.sum(np.array(null_aucs) >= mean_obs) + 1) / (_N_PERM + 1))
        sig    = '*' if perm_p < 0.05 else ''
        print(f'{clf_name:<22} {mean_obs:>10.3f} {sd_obs:>6.3f}  {perm_p:.3f}   {sig}')
        clf_results.append(ClassifierResult(clf_name, mean_obs, sd_obs, perm_p, sig))

        if mean_obs > best_mean_auc:
            best_mean_auc  = mean_obs
            best_clf_name  = clf_name
            best_pipe      = pipe

        # Refit on all data to extract stable feature importances
        if clf_name == 'Random Forest':
            pipe.fit(X, y)
            rf_importances = pipe.named_steps['clf'].feature_importances_

    print('* p < 0.05 (permutation test)')
    print(f'\nNote: with n={n_total} and class split ({n_pos}:{n_total - n_pos}), AUC estimates')
    print('  have wide variance; interpret with caution.')

    # Refit winning model on training data, then apply to evaluation group
    assert best_pipe is not None
    best_pipe.fit(X, y)
    eval_feat  = _build_features(df_eval)
    eval_y     = _compute_y(df_eval, target)
    eval_valid = eval_feat.notna().all(axis=1) & eval_y.notna()
    X_eval     = eval_feat[eval_valid].to_numpy()
    y_eval     = eval_y[eval_valid]
    ids_eval   = df_eval.loc[eval_valid, 'Row']
    prob_eval  = best_pipe.predict_proba(X_eval)[:, 1]

    print(f'\n{eval_label} group predictions — winning model: {best_clf_name} (AUC = {best_mean_auc:.3f})')
    print(f'  Target: {target.name}')
    print(f'  n = {len(y_eval)} {eval_label.lower()} participants')
    print()
    print(f"{'ID':>5}  {'Actual':>8}  {'Pred. prob':>10}  {'Predicted':>9}")
    print('-' * 40)
    n_pred_pos = int((prob_eval >= 0.5).sum())
    for pid, actual, prob in zip(ids_eval, y_eval, prob_eval):
        predicted = 'Yes' if prob >= 0.5 else 'No'
        actual_str = 'Yes' if actual == 1 else 'No'
        print(f'{int(pid):>5}  {actual_str:>8}  {prob:>10.3f}  {predicted:>9}')
    print()
    print(f'  Predicted positive: {n_pred_pos}/{len(y_eval)} ({n_pred_pos / len(y_eval):.0%})')

    out_dir.mkdir(parents=True, exist_ok=True)
    _plot_roc_curves(all_fpr, all_tpr, cv_aucs, out_dir, len(feature_names), target)
    if rf_importances is not None:
        _plot_feature_importance(rf_importances, feature_names, out_dir, target)
    print(f'\nPlots saved to {out_dir}/')

    _write_docx(clf_results, n_total, n_pos, len(feature_names), target, docx_path)


def _write_docx(
    results: list[ClassifierResult],
    n_total: int,
    n_pos: int,
    n_features: int,
    target: Target,
    out_path: Path,
) -> None:
    """Save classifier performance table → docx."""
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['Classifier', 'Mean AUC', 'SD', 'Perm. p', 'Sig.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_col_widths(table, [4.5, 2.5, 2.0, 2.5, 1.2])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    for r in results:
        row = table.add_row()
        for cell in row.cells:
            clear_cell(cell)
        for cell, val in zip(row.cells,
                              [r.name, f'{r.mean_auc:.3f}', f'{r.sd_auc:.3f}',
                               f'{r.perm_p:.3f}', r.sig]):
            add_run(cell.paragraphs[0], val, FONT_BODY)

    legend = (
        f'ML prediction of {target.name.lower()} (n={n_total}; '
        f'positive: {n_pos}, negative: {n_total - n_pos}). '
        f'Predictors: {n_features} features (biomarker T0 values, sex, chronological age). '
        f'Stratified 5-fold cross-validation; metric: AUC-ROC (mean ± SD across folds). '
        f'Permutation p from {_N_PERM} label-shuffled runs (+1/(N+1) continuity correction). '
        f'* p < 0.05.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(str(out_path))
    print(f'Saved {out_path}')


def _plot_roc_curves(
    all_fpr: dict[str, list[np.ndarray]],
    all_tpr: dict[str, list[np.ndarray]],
    cv_aucs: dict[str, list[float]],
    out_dir: Path,
    n_features: int,
    target: Target,
) -> None:
    """Mean ROC curve across CV folds for each classifier, with AUC ± SD in legend."""
    fig, ax = plt.subplots(figsize=(6, 6))
    mean_fpr = np.linspace(0, 1, 100)

    for clf_name, fpr_list in all_fpr.items():
        tpr_list = all_tpr[clf_name]
        # np.interp: linearly interpolate each fold's TPR onto the shared FPR grid
        interp_tprs = [np.interp(mean_fpr, fpr, tpr) for fpr, tpr in zip(fpr_list, tpr_list)]
        mean_tpr = np.mean(interp_tprs, axis=0)
        mean_auc = float(np.mean(cv_aucs[clf_name]))
        std_auc  = float(np.std(cv_aucs[clf_name]))
        ax.plot(mean_fpr, mean_tpr, color=_CLF_COLORS[clf_name], linewidth=2,
                label=f'{clf_name} (AUC = {mean_auc:.3f} ± {std_auc:.3f})')

    ax.plot([0, 1], [0, 1], 'k--', linewidth=0.8, label='Chance')
    ax.set_xlabel('False Positive Rate', fontsize=10)
    ax.set_ylabel('True Positive Rate', fontsize=10)
    ax.set_title(
        f'ROC Curves – {target.name} Prediction\n'
        f'(Stratified 5-Fold CV, {n_features} T0 Predictors)',
        fontsize=10,
    )
    ax.legend(fontsize=8, loc='lower right')
    fig.tight_layout()
    fig.savefig(out_dir / 'roc_curves.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {out_dir}/roc_curves.png')


def _plot_feature_importance(
    importances: np.ndarray,
    names: list[str],
    out_dir: Path,
    target: Target,
) -> None:
    """Horizontal bar chart of Random Forest feature importances, sorted ascending for readability."""
    idx          = np.argsort(importances)   # ascending → bottom-to-top
    sorted_names = [names[i] for i in idx]
    sorted_imp   = importances[idx]

    fig, ax = plt.subplots(figsize=(7, max(4, len(names) * 0.4 + 1)))
    y_pos = np.arange(len(sorted_names))
    ax.barh(y_pos, sorted_imp * 100, color='#0072B2', alpha=0.85)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(sorted_names, fontsize=8)
    ax.set_xlabel('Feature importance (%)', fontsize=10)
    ax.set_title(
        f'Random Forest Feature Importances\n'
        f'({target.name}, {len(names)} T0 Predictors)',
        fontsize=10,
    )
    fig.tight_layout()
    fig.savefig(out_dir / 'feature_importance.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {out_dir}/feature_importance.png')
