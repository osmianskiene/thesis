from pathlib import Path
from typing import cast

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from scipy.stats import mannwhitneyu
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler

from biomarkers import ANALYSIS_BIOMARKERS, AGE_ACCEL_BIOMARKERS
from commands.univariate import _cohens_d, _effect_label
from data import load_composites
from docx_utils import C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

_COLOR_TREAT = '#0072B2'
_COLOR_CTRL  = '#E69F00'

_OUT_DIR = Path('output/plots_composite_scores')

# Names and "beneficial direction" sign for each composite:
# +1 means higher raw score = beneficial (flip so positive = improvement)
# -1 means lower raw score = beneficial (no flip)
_AGE_ACCEL_DELTA_COLS = [bm.delta_col for bm in AGE_ACCEL_BIOMARKERS]


def composite_scores() -> None:
    """Three composite ageing scores → stdout + output/a2_composite_scores.docx + plots."""
    df = load_composites()

    delta_cols = [bm.delta_col for bm in ANALYSIS_BIOMARKERS]
    names      = [bm.name      for bm in ANALYSIS_BIOMARKERS]
    directions = [bm.direction for bm in ANALYSIS_BIOMARKERS]

    # Drop rows with any missing Δ; keep group labels aligned
    X_raw  = df[delta_cols].dropna()
    groups = df.loc[X_raw.index, 'Treatment']

    # ── Sign-flip so that positive z-score always means "beneficial change" ──
    # For 'down' biomarkers: improvement = negative Δ → multiply by -1
    # For 'up' biomarkers:   improvement = positive Δ → keep as-is
    sign = np.array([-1.0 if d == 'down' else 1.0 for d in directions])

    # Z-score raw Δ values (mean=0, SD=1 across all participants)
    scaler   = StandardScaler()
    X_scaled = scaler.fit_transform(X_raw)  # shape: (n_samples, 20)

    # Sign-flip after standardization so positive z = improvement
    X_benefit = X_scaled * sign[np.newaxis, :]

    # ── Score 1: Equal-weighted ───────────────────────────────────────────────
    # Mean of 20 sign-corrected z-scores; positive = net beneficial shift
    score_equal = pd.Series(X_benefit.mean(axis=1), index=X_raw.index, name='Equal-weighted')

    # ── Score 2: PC1-weighted ─────────────────────────────────────────────────
    # Project onto PC1 of the original (unsigned) z-scores, then orient so that
    # the treatment group has a positive mean (PC1 sign is arbitrary).
    pca    = PCA(n_components=1)
    pc1    = pca.fit_transform(X_scaled).ravel()  # shape: (n_samples,)
    t_mask = groups.values == 1
    # Orient so positive = treatment direction
    if pc1[t_mask].mean() < pc1[~t_mask].mean():
        pc1 = -pc1
    score_pc1 = pd.Series(pc1, index=X_raw.index, name='PC1-weighted')

    # ── Score 3: Ageing-specific ──────────────────────────────────────────────
    # Mean of the four age-acceleration Δ values (already in age-acceleration units).
    # Multiply by -1 so positive = improvement (lower age-acceleration = better).
    age_cols = [c for c in _AGE_ACCEL_DELTA_COLS if c in df.columns]
    age_raw  = df.loc[X_raw.index, age_cols].mean(axis=1)  # NaN if all missing
    score_age = (-age_raw).rename('Ageing-specific')  # flip so positive = improvement

    scores = {
        'Equal-weighted':   score_equal,
        'PC1-weighted':     score_pc1,
        'Ageing-specific':  score_age,
    }

    # ── Print results ─────────────────────────────────────────────────────────
    print("Composite Ageing Scores – Group Comparison (Mann-Whitney U)")
    print("=" * 80)
    print(f"{'Score':<22} {'Treat mean':>11} {'Ctrl mean':>10} {'U':>10} {'p':>8} {'d':>7} {'Effect':>8}  Sig.")
    print("-" * 80)

    rows: list[tuple] = []
    for label, s in scores.items():
        t_vals = s[groups == 1].dropna()
        c_vals = s[groups == 0].dropna()
        u_stat, p = mannwhitneyu(t_vals, c_vals, alternative='two-sided')
        sig = '*' if p < 0.05 else ''
        # rank-biserial correlation as effect size: r = 1 - 2U / (n_t * n_c)
        r_rb = 1 - 2 * u_stat / (len(t_vals) * len(c_vals))
        d    = _cohens_d(t_vals, c_vals)
        rows.append((label, t_vals.mean(), c_vals.mean(), u_stat, p, r_rb, d))
        print(f"{label:<22} {t_vals.mean():>11.3f} {c_vals.mean():>10.3f} "
              f"{u_stat:>10.1f} {p:>8.4f} {d:>7.3f} {_effect_label(d):>8}  {sig}")

    print("* p < 0.05")
    print("\nNote: positive score = net beneficial direction for each composite.")
    print(f"  Equal-weighted:  mean of {len(delta_cols)} sign-corrected standardized Δ values")
    print(f"  PC1-weighted:    projection onto PC1 ({pca.explained_variance_ratio_[0]*100:.1f}% variance),")
    print(f"                   oriented so positive = treatment direction")
    print(f"  Ageing-specific: mean of {len(age_cols)} age-acceleration Δ values (sign-flipped, positive = improvement)")

    _write_docx(rows, pca.explained_variance_ratio_[0])
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    _plot_boxplots(scores, groups)
    print(f"\nPlots saved to {_OUT_DIR}/")


def _write_docx(rows: list[tuple], pc1_var: float) -> None:
    """Save composite scores group comparison → output/a2_composite_scores.docx."""
    out = Path('output/a2_composite_scores.docx')
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['Composite score', 'Treat mean', 'Ctrl mean', 'U statistic', 'p', 'd', 'r_rb', 'Sig.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_col_widths(table, [3.6, 2.2, 2.2, 2.4, 2.0, 1.6, 1.8, 1.2])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    for label, t_mean, c_mean, u_stat, p, r_rb, d in rows:
        sig = '*' if p < 0.05 else ''
        data = [label, f'{t_mean:.3f}', f'{c_mean:.3f}',
                f'{u_stat:.1f}', f'{p:.4f}', f'{d:.3f}', f'{r_rb:.3f}', sig]
        row = table.add_row()
        for cell, val in zip(row.cells, data):
            clear_cell(cell)
            add_run(cell.paragraphs[0], val, FONT_BODY)

    legend = (
        f'Mann-Whitney U test (two-sided), Treatment vs Control. '
        f'Positive score = net beneficial direction. '
        f'Equal-weighted: mean of 20 sign-corrected standardized Δ values. '
        f'PC1-weighted: projection onto PC1 ({pc1_var*100:.1f}% variance), oriented toward treatment. '
        f'Ageing-specific: mean of 4 age-acceleration Δ values (sign-flipped; positive = improvement). '
        f'r_rb = rank-biserial correlation (effect size). * p < 0.05.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(out)
    print(f'Saved {out}')


def _plot_boxplots(scores: dict[str, pd.Series], groups: pd.Series) -> None:
    """Side-by-side boxplots for each composite score, Treatment vs Control."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 5), sharey=False)

    for ax, (label, s) in zip(axes, scores.items()):
        t_vals = s[groups == 1].dropna().values
        c_vals = s[groups == 0].dropna().values

        bp = ax.boxplot(
            [t_vals, c_vals],
            labels=['Treatment', 'Control'],
            patch_artist=True,
            showmeans=True,
            medianprops=dict(color='black', linewidth=1.5),           # median: solid black line
            meanprops=dict(marker='D', markerfacecolor='white',
                           markeredgecolor='black', markersize=6),    # mean: white diamond
            whiskerprops=dict(linewidth=1),
            capprops=dict(linewidth=1),
        )
        bp['boxes'][0].set_facecolor(_COLOR_TREAT + '99')  # hex alpha
        bp['boxes'][1].set_facecolor(_COLOR_CTRL  + '99')

        # Overlay individual points with jitter
        rng = np.random.default_rng(42)
        for i, (vals, color) in enumerate([(t_vals, _COLOR_TREAT), (c_vals, _COLOR_CTRL)], 1):
            jitter = rng.uniform(-0.12, 0.12, size=len(vals))
            ax.scatter(np.full(len(vals), i) + jitter, vals,
                       color=color, alpha=0.55, s=22, zorder=3)

        # Horizontal reference at 0 (no net change)
        ax.axhline(0, color='grey', linewidth=0.7, linestyle='--')
        ax.set_title(label, fontsize=10, fontweight='bold')
        ax.set_ylabel('Score (positive = beneficial)' if ax == axes[0] else '')
        ax.tick_params(axis='x', labelsize=9)

    # Shared legend: median line + mean diamond
    from matplotlib.lines import Line2D
    from matplotlib.patches import Patch
    legend_handles = [
        Line2D([0], [0], color='black', linewidth=1.5, label='Median'),
        Line2D([0], [0], marker='D', color='w', markerfacecolor='white',
               markeredgecolor='black', markersize=6, label='Mean'),
    ]
    axes[-1].legend(handles=legend_handles, fontsize=8, loc='upper right')

    fig.suptitle('Composite Ageing Scores – Treatment vs Control', fontsize=11, y=1.01)
    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'composite_scores_boxplots.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {_OUT_DIR}/composite_scores_boxplots.png')
