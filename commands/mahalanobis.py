from pathlib import Path
from typing import cast, NamedTuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
import pingouin as pg

from biomarkers import ANALYSIS_BIOMARKERS
from data import load_composites
from docx_utils import C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

# Okabe-Ito: blue = Treatment, orange = Control
_COLOR_TREAT = '#0072B2'
_COLOR_CTRL  = '#E69F00'
_COLOR_NULL  = '#56B4E9'  # light blue for null distribution bars
_COLOR_OBS   = '#D55E00'  # vermilion for observed statistic line

_OUT_DIR = Path('output/plots_mahalanobis')

_N_PERM = 5_000


class MahalanobisResult(NamedTuple):
    # key quantities for the between-centroid test
    D2_obs: float           # observed squared Mahalanobis distance
    T2_obs: float           # Hotelling's T² statistic
    F_stat: float           # F-approximation statistic (nan if df2 ≤ 0)
    F_df1: int              # numerator degrees of freedom
    F_df2: int              # denominator degrees of freedom
    F_p: float              # parametric p-value (nan if not available)
    perm_p: float           # permutation p-value
    null_mean: float        # mean of permutation null distribution
    null_95: float          # 95th percentile of null distribution
    n_treat: int
    n_ctrl: int
    n_biomarkers: int
    inv_method: str         # 'exact inverse' or 'Moore-Penrose pseudoinverse'
    cond_num: float         # condition number of pooled covariance matrix
    # centroid-assignment accuracy
    accuracy: float
    n_correct_treat: int
    n_correct_ctrl: int
    # per-observation distances (all participants)
    dist_to_treat: np.ndarray
    dist_to_ctrl: np.ndarray
    treat_mask: np.ndarray  # boolean, True = treatment
    # permutation null distribution (for plotting)
    null_dist: np.ndarray
    # correlation filter metadata (empty lists when filter not applied)
    kept_names: list[str]
    dropped_names: list[str]
    corr_threshold: float | None  # None = filter not applied


def _corr_filter(X: np.ndarray, names: list[str], threshold: float) -> tuple[np.ndarray, list[str], list[str]]:
    """Drop one biomarker from every pair whose |Pearson r| exceeds threshold.

    Greedy: iterate columns left-to-right; when a pair exceeds the threshold,
    the later column is dropped (earlier = listed first in ANALYSIS_BIOMARKERS).
    Correlation is computed on all participants combined so group labels play
    no role in the selection — this makes the filter outcome-blind.
    """
    corr = np.corrcoef(X.T)  # shape: (p, p)
    p = X.shape[1]
    keep = list(range(p))   # indices of columns retained so far

    dropped_idx: list[int] = []
    for i in range(len(keep)):
        for j in range(i + 1, len(keep)):
            if keep[j] in dropped_idx:
                continue
            if abs(corr[keep[i], keep[j]]) > threshold:
                dropped_idx.append(keep[j])

    kept_idx    = [i for i in range(p) if i not in dropped_idx]
    kept_names  = [names[i] for i in kept_idx]
    dropped_names = [names[i] for i in dropped_idx]
    return X[:, kept_idx], kept_names, dropped_names


def mahalanobis_analysis(corr_threshold: float | None = None) -> None:
    """Mahalanobis D² between group centroids with 5,000-permutation test → stdout + plots + docx."""
    df = load_composites()

    delta_cols = [bm.delta_col for bm in ANALYSIS_BIOMARKERS]
    names      = [bm.name      for bm in ANALYSIS_BIOMARKERS]

    X_raw  = df[delta_cols].dropna()
    groups = df.loc[X_raw.index, 'Treatment']

    # Standardize Δ values (mean=0, SD=1) so scale differences don't dominate
    from sklearn.preprocessing import StandardScaler
    scaler   = StandardScaler()
    X_scaled = scaler.fit_transform(X_raw)  # shape: (n_samples, p)

    dropped_names: list[str] = []
    kept_names = names
    if corr_threshold is not None:
        X_scaled, kept_names, dropped_names = _corr_filter(X_scaled, names, corr_threshold)
        if dropped_names:
            print(f'Correlation filter (|r| > {corr_threshold}): dropped {len(dropped_names)} biomarker(s):')
            for name in dropped_names:
                print(f'  - {name}')
        else:
            print(f'Correlation filter (|r| > {corr_threshold}): no pairs exceeded threshold.')

    result = _compute(X_scaled, groups.to_numpy(), kept_names, dropped_names, corr_threshold)

    suffix  = '_corr_filter' if corr_threshold is not None else ''
    out_dir = Path(f'output/plots_mahalanobis{suffix}')

    _print_results(result)
    _write_docx(result, suffix)
    out_dir.mkdir(parents=True, exist_ok=True)
    _plot_permutation(result, out_dir)
    _plot_centroid_distances(result, out_dir)
    print(f'\nPlots saved to {out_dir}/')


def _compute(X: np.ndarray, labels: np.ndarray,
             kept_names: list[str], dropped_names: list[str],
             corr_threshold: float | None) -> MahalanobisResult:
    """Compute Mahalanobis D², Hotelling T², permutation test, and per-obs distances."""
    treat_mask = labels == 1
    ctrl_mask  = labels == 0

    X_t = X[treat_mask]
    X_c = X[ctrl_mask]
    n_t, p = X_t.shape
    n_c    = X_c.shape[0]

    mu_t = X_t.mean(axis=0)
    mu_c = X_c.mean(axis=0)
    diff = mu_t - mu_c

    # Pooled within-group covariance: weighted average of each group's covariance
    S_t    = np.cov(X_t.T)
    S_c    = np.cov(X_c.T)
    S_pool = ((n_t - 1) * S_t + (n_c - 1) * S_c) / (n_t + n_c - 2)

    cond_num = float(np.linalg.cond(S_pool))
    try:
        S_inv      = np.linalg.inv(S_pool)
        inv_method = 'exact inverse'
    except np.linalg.LinAlgError:
        S_inv      = np.linalg.pinv(S_pool)  # pseudoinverse as fallback
        inv_method = 'Moore-Penrose pseudoinverse'

    # D²_M = (µ_t − µ_c)ᵀ S_pool⁻¹ (µ_t − µ_c)
    D2_obs = float(diff @ S_inv @ diff)

    # Hotelling's T² and F-approximation via pingouin (battle-tested implementation)
    ht = pg.multivariate_ttest(X_t, X_c)
    T2_obs = float(ht['T2'].iloc[0])
    F_stat = float(ht['F'].iloc[0])
    F_df1  = int(ht['df1'].iloc[0])
    F_df2  = int(ht['df2'].iloc[0])
    F_p    = float(ht['pval'].iloc[0])

    # Permutation test: shuffle group labels, recompute D² against fixed S_pool
    rng     = np.random.default_rng(42)
    null_d2 = np.empty(_N_PERM)
    for k in range(_N_PERM):
        perm  = rng.permutation(labels)
        d_perm = X[perm == 1].mean(axis=0) - X[perm == 0].mean(axis=0)
        null_d2[k] = float(d_perm @ S_inv @ d_perm)

    # +1 / (N+1) continuity correction avoids p = 0
    perm_p = float((np.sum(null_d2 >= D2_obs) + 1) / (_N_PERM + 1))

    # Per-observation Mahalanobis distance to each centroid
    # d_M(x, µ) = sqrt((x − µ)ᵀ S_pool⁻¹ (x − µ))
    # einsum 'ij,jk,ik->i' computes the quadratic form for each row simultaneously
    def _row_dist(Z: np.ndarray, mu: np.ndarray) -> np.ndarray:
        d = Z - mu
        return np.sqrt(np.einsum('ij,jk,ik->i', d, S_inv, d))

    dist_to_treat = _row_dist(X, mu_t)
    dist_to_ctrl  = _row_dist(X, mu_c)

    # Nearest-centroid classification: assign each observation to the closer centroid
    predicted_treat = dist_to_treat < dist_to_ctrl
    accuracy        = float((predicted_treat == treat_mask).mean())
    n_correct_treat = int((predicted_treat & treat_mask).sum())
    n_correct_ctrl  = int((~predicted_treat & ctrl_mask).sum())

    return MahalanobisResult(
        D2_obs=D2_obs, T2_obs=T2_obs,
        F_stat=F_stat, F_df1=F_df1, F_df2=F_df2, F_p=F_p,
        perm_p=perm_p,
        null_mean=float(null_d2.mean()), null_95=float(np.percentile(null_d2, 95)),
        n_treat=n_t, n_ctrl=n_c, n_biomarkers=p,
        inv_method=inv_method, cond_num=cond_num,
        accuracy=accuracy,
        n_correct_treat=n_correct_treat, n_correct_ctrl=n_correct_ctrl,
        dist_to_treat=dist_to_treat, dist_to_ctrl=dist_to_ctrl,
        treat_mask=treat_mask,
        null_dist=null_d2,
        kept_names=kept_names,
        dropped_names=dropped_names,
        corr_threshold=corr_threshold,
    )


def _print_results(r: MahalanobisResult) -> None:
    print("Mahalanobis Distance – Multivariate Group Separation")
    print("=" * 60)
    print(f"  Biomarkers (p): {r.n_biomarkers}"
          + (f" (filtered from 20, |r| > {r.corr_threshold})" if r.corr_threshold else ""))
    if r.dropped_names:
        print(f"  Dropped:        {', '.join(r.dropped_names)}")
    print(f"  Kept:           {', '.join(r.kept_names)}")
    print(f"  Treatment n:    {r.n_treat}")
    print(f"  Control n:      {r.n_ctrl}")
    print(f"  Covariance:     pooled within-group ({r.inv_method})")
    print(f"  Condition no.:  {r.cond_num:.1f}")
    print()
    print(f"  Observed D²_M  = {r.D2_obs:.4f}")
    print(f"  Hotelling T²   = {r.T2_obs:.4f}")
    if not np.isnan(r.F_p):
        print(f"  F-approx:      F({r.F_df1},{r.F_df2}) = {r.F_stat:.3f}, p = {r.F_p:.4f}  "
              f"[parametric; assumes multivariate normality]")
    else:
        print(f"  F-approx:      not available (n − p − 1 ≤ 0)")
    print()
    print(f"  Permutation test ({_N_PERM:,} permutations, fixed S_pool):")
    print(f"    p-value = {r.perm_p:.4f}  "
          f"({'SIGNIFICANT' if r.perm_p < 0.05 else 'not significant'}, α=0.05)")
    print(f"    Null mean = {r.null_mean:.4f},  95th pct = {r.null_95:.4f}")
    print()
    print(f"  Nearest-centroid assignment accuracy: {r.accuracy*100:.1f}%")
    print(f"    Treatment → assigned Treatment: {r.n_correct_treat}/{r.n_treat}")
    print(f"    Control   → assigned Control:  {r.n_correct_ctrl}/{r.n_ctrl}")


def _write_docx(r: MahalanobisResult, suffix: str = '') -> None:
    """Save Mahalanobis results summary table → output/a2_mahalanobis[suffix].docx."""
    out = Path(f'output/a2_mahalanobis{suffix}.docx')
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['Statistic', 'Value']
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_col_widths(table, [7.0, 5.5])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    def _row(label: str, value: str) -> None:
        row = table.add_row()
        clear_cell(row.cells[0])
        clear_cell(row.cells[1])
        add_run(row.cells[0].paragraphs[0], label, FONT_BODY)
        add_run(row.cells[1].paragraphs[0], value, FONT_BODY)

    _row('Biomarkers (p)', str(r.n_biomarkers)
         + (f' (filtered from 20, |r| > {r.corr_threshold})' if r.corr_threshold else ''))
    if r.dropped_names:
        _row('Dropped (high correlation)', ', '.join(r.dropped_names))
    _row('Treatment n', str(r.n_treat))
    _row('Control n', str(r.n_ctrl))
    _row('Covariance matrix', f'Pooled within-group ({r.inv_method})')
    _row('Condition number', f'{r.cond_num:.1f}')
    _row('Observed D²_M', f'{r.D2_obs:.4f}')
    _row("Hotelling's T²", f'{r.T2_obs:.4f}')
    if not np.isnan(r.F_p):
        _row(f'F-approximation F({r.F_df1},{r.F_df2})', f'{r.F_stat:.3f}')
        _row('Parametric p (F)', f'{r.F_p:.4f}')
    else:
        _row('F-approximation', 'Not available (n − p − 1 ≤ 0)')
    _row(f'Permutation p ({_N_PERM:,} permutations)', f'{r.perm_p:.4f}')
    _row('Permutation null mean', f'{r.null_mean:.4f}')
    _row('Permutation null 95th pct', f'{r.null_95:.4f}')
    _row('Nearest-centroid accuracy', f'{r.accuracy*100:.1f}%')
    _row('Treatment correctly assigned', f'{r.n_correct_treat}/{r.n_treat}')
    _row('Control correctly assigned', f'{r.n_correct_ctrl}/{r.n_ctrl}')

    sig = 'Yes (p < 0.05)' if r.perm_p < 0.05 else 'No (p ≥ 0.05)'
    _row('Significant multivariate separation', sig)

    legend = (
        f'Squared Mahalanobis distance (D²_M) between group centroids computed using '
        f'the pooled within-group covariance matrix on {r.n_biomarkers} standardized '
        f'biomarker Δ values. Permutation test: group labels shuffled {_N_PERM:,} times '
        f'with S_pool held fixed. Parametric F-approximation assumes multivariate '
        f'normality (Hotelling\'s T²). Nearest-centroid accuracy: proportion of '
        f'participants assigned to their true group by minimum Mahalanobis distance.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(str(out))
    print(f'Saved {out}')


def _plot_permutation(r: MahalanobisResult, out_dir: Path = _OUT_DIR) -> None:
    """Histogram of permutation null distribution with observed D² and 95th pct marked."""
    fig, ax = plt.subplots(figsize=(7, 5))

    ax.hist(r.null_dist, bins=60, color=_COLOR_NULL, alpha=0.75, density=True,
            label=f'Null distribution ({_N_PERM:,} permutations)')
    ax.axvline(r.D2_obs, color=_COLOR_OBS, linewidth=2,
               label=f'Observed $D^2_M$ = {r.D2_obs:.2f}')
    ax.axvline(r.null_95, color='#555555', linewidth=1.2, linestyle='--',
               label=f'95th percentile = {r.null_95:.2f}')

    sig_text = f'p = {r.perm_p:.4f}' if r.perm_p >= 0.001 else 'p < 0.001'
    ax.set_xlabel('$D^2_M$ (squared Mahalanobis distance)', fontsize=10)
    ax.set_ylabel('Density', fontsize=10)
    ax.set_title(
        f'Permutation Null Distribution – Mahalanobis $D^2_M$\n'
        f'{sig_text}  ({"significant" if r.perm_p < 0.05 else "not significant"}, α = 0.05)',
        fontsize=10,
    )
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(out_dir / 'permutation_null.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {out_dir}/permutation_null.png')


def _plot_centroid_distances(r: MahalanobisResult, out_dir: Path = _OUT_DIR) -> None:
    """Scatter of per-observation distances to each centroid, coloured by true group."""
    fig, ax = plt.subplots(figsize=(6, 6))

    ctrl_mask = ~r.treat_mask

    ax.scatter(r.dist_to_treat[r.treat_mask], r.dist_to_ctrl[r.treat_mask],
               color=_COLOR_TREAT, alpha=0.65, s=50, edgecolors='white',
               linewidth=0.4, label='Treatment', zorder=3)
    ax.scatter(r.dist_to_treat[ctrl_mask], r.dist_to_ctrl[ctrl_mask],
               color=_COLOR_CTRL, alpha=0.65, s=50, edgecolors='white',
               linewidth=0.4, label='Control', zorder=3)

    # Diagonal: equal distance to both centroids → points above = nearer to Control
    lim = max(r.dist_to_treat.max(), r.dist_to_ctrl.max()) * 1.07
    ax.plot([0, lim], [0, lim], color='#888888', linewidth=0.8, linestyle='--',
            label='Equal distance')
    ax.set_xlim(0, lim)
    ax.set_ylim(0, lim)

    ax.set_xlabel('$D_M$ to Treatment centroid', fontsize=10)
    ax.set_ylabel('$D_M$ to Control centroid', fontsize=10)
    ax.set_title(
        f'Per-Observation Mahalanobis Distances\n'
        f'Nearest-centroid accuracy = {r.accuracy*100:.0f}%',
        fontsize=10,
    )
    ax.legend(fontsize=8, loc='upper left')
    ax.set_aspect('equal')

    fig.tight_layout()
    fig.savefig(out_dir / 'centroid_distances.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {out_dir}/centroid_distances.png')

    _plot_combined(r, out_dir)


def _plot_combined(r: MahalanobisResult, out_dir: Path = _OUT_DIR) -> None:
    """Combined figure: centroid distances (left) + permutation null (right)."""
    fig, (ax_cd, ax_pn) = plt.subplots(1, 2, figsize=(13, 6))

    # ── Left: centroid distances ──────────────────────────────────────────────
    ctrl_mask = ~r.treat_mask
    ax_cd.scatter(r.dist_to_treat[r.treat_mask], r.dist_to_ctrl[r.treat_mask],
                  color=_COLOR_TREAT, alpha=0.65, s=50, edgecolors='white',
                  linewidth=0.4, label='Treatment', zorder=3)
    ax_cd.scatter(r.dist_to_treat[ctrl_mask], r.dist_to_ctrl[ctrl_mask],
                  color=_COLOR_CTRL, alpha=0.65, s=50, edgecolors='white',
                  linewidth=0.4, label='Control', zorder=3)
    lim = max(r.dist_to_treat.max(), r.dist_to_ctrl.max()) * 1.07
    ax_cd.plot([0, lim], [0, lim], color='#888888', linewidth=0.8, linestyle='--',
               label='Equal distance')
    ax_cd.set_xlim(0, lim)
    ax_cd.set_ylim(0, lim)
    ax_cd.set_xlabel('$D_M$ to Treatment centroid', fontsize=10)
    ax_cd.set_ylabel('$D_M$ to Control centroid', fontsize=10)
    ax_cd.set_title(
        f'Per-Observation Mahalanobis Distances\n'
        f'Nearest-centroid accuracy = {r.accuracy*100:.0f}%',
        fontsize=10,
    )
    ax_cd.legend(fontsize=8, loc='upper left')
    ax_cd.set_aspect('equal')

    # ── Right: permutation null distribution ──────────────────────────────────
    ax_pn.hist(r.null_dist, bins=60, color=_COLOR_NULL, alpha=0.75, density=True,
               label=f'Null distribution ({_N_PERM:,} permutations)')
    ax_pn.axvline(r.D2_obs, color=_COLOR_OBS, linewidth=2,
                  label=f'Observed $D^2_M$ = {r.D2_obs:.2f}')
    ax_pn.axvline(r.null_95, color='#555555', linewidth=1.2, linestyle='--',
                  label=f'95th percentile = {r.null_95:.2f}')
    sig_text = f'p = {r.perm_p:.4f}' if r.perm_p >= 0.001 else 'p < 0.001'
    ax_pn.set_xlabel('$D^2_M$ (squared Mahalanobis distance)', fontsize=10)
    ax_pn.set_ylabel('Density', fontsize=10)
    ax_pn.set_title(
        f'Permutation Null Distribution – Mahalanobis D²\n'
        f'{sig_text}  ({"significant" if r.perm_p < 0.05 else "not significant"}, α = 0.05)',
        fontsize=10,
    )
    ax_pn.legend(fontsize=8)

    fig.tight_layout()
    fig.savefig(out_dir / 'combined.png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    print(f'Saved {out_dir}/combined.png')
