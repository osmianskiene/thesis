from pathlib import Path
from typing import cast

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.styles.style import _ParagraphStyle
from scipy.stats import mannwhitneyu
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler

from biomarkers import ANALYSIS_BIOMARKERS
from data import load_composites
from docx_utils import C_BLACK, C_GRAY, FONT_BODY, FONT_LEGEND, add_run, clear_cell, set_col_widths

# Okabe-Ito: blue = Treatment, orange = Control
_COLOR_TREAT = '#0072B2'
_COLOR_CTRL  = '#E69F00'

_OUT_DIR = Path('output/plots_pca')


def pca_analysis() -> None:
    """PCA on 20 standardized biomarker Δ values → stdout."""
    df = load_composites()

    delta_cols = [bm.delta_col for bm in ANALYSIS_BIOMARKERS]
    names      = [bm.name      for bm in ANALYSIS_BIOMARKERS]

    # Drop rows with any missing Δ; keep group labels aligned
    X_raw  = df[delta_cols].dropna()
    groups = df.loc[X_raw.index, 'Treatment']

    # Standardize to z-scores (mean=0, SD=1) so units don't dominate
    scaler   = StandardScaler()
    X_scaled = scaler.fit_transform(X_raw)  # shape: (n_samples, 20)

    # Full PCA — all components, then we decide how many to retain
    pca    = PCA()
    scores = pca.fit_transform(X_scaled)  # shape: (n_samples, 20)

    eigenvalues = pca.explained_variance_         # variance (eigenvalue) per PC
    var_ratio   = pca.explained_variance_ratio_   # fraction of total variance
    cum_var     = np.cumsum(var_ratio)

    # Retention criteria
    n_kaiser = int((eigenvalues > 1).sum())           # Kaiser: eigenvalue > 1
    n_80pct  = int(np.searchsorted(cum_var, 0.80)) + 1  # ≥80% cumulative variance
    n_retain = max(n_kaiser, n_80pct)

    # ── Variance explained table ──────────────────────────────────────────────
    print("PCA – Variance Explained")
    print("=" * 58)
    print(f"{'PC':<5} {'Eigenvalue':>11} {'Var%':>7} {'Cum%':>8}  Kaiser")
    print("-" * 58)
    for i, (ev, vr, cv) in enumerate(zip(eigenvalues, var_ratio, cum_var)):
        kaiser = "✓" if ev > 1 else ""
        flag   = "  ← 80% threshold" if i + 1 == n_80pct else ""
        print(f"PC{i+1:<4} {ev:>11.3f} {vr*100:>6.1f}% {cv*100:>7.1f}%  {kaiser}{flag}")

    print(f"\nKaiser criterion (eigenvalue > 1): {n_kaiser} component(s)")
    print(f"Cumulative variance ≥80%:          {n_80pct} component(s)  "
          f"({cum_var[n_80pct-1]*100:.1f}% retained)")
    print(f"Components retained for analysis:  {n_retain}")

    # ── ASCII scree plot ──────────────────────────────────────────────────────
    print("\nScree Plot")
    print("-" * 50)
    bar_max = 35
    for i in range(len(eigenvalues)):
        ev    = eigenvalues[i]
        bar   = int(ev / eigenvalues[0] * bar_max)
        above = "●" if ev > 1 else "○"  # filled = above Kaiser threshold
        print(f"PC{i+1:<3} {above} {'█' * bar} {ev:.3f}")
    print(f"       (● = eigenvalue > 1;  Kaiser threshold at 1.000)")

    # ── PC loadings for retained components ──────────────────────────────────
    loadings = pca.components_  # shape: (n_components, n_features)

    print(f"\nPC Loadings  (retained: PC1–PC{n_retain})")
    print("=" * (28 + 9 * n_retain))
    header = f"{'Biomarker':<28}" + "".join(f"{'PC'+str(i+1):>9}" for i in range(n_retain))
    print(header)
    print("-" * (28 + 9 * n_retain))

    # Sort by descending |loading| on PC1 so the dominant biomarkers appear first
    sort_idx = np.argsort(np.abs(loadings[0]))[::-1]
    for j in sort_idx:
        row = f"{names[j]:<28}"
        for i in range(n_retain):
            row += f"{loadings[i, j]:>9.3f}"
        print(row)

    # ── Per-PC: which biomarkers drive it most ───────────────────────────────
    print(f"\nPer-PC Main Biomarkers  (|loading| ≥ 0.20, sorted by |loading|)")
    print("=" * 70)
    for i in range(n_retain):
        bm_loads = sorted(
            ((abs(loadings[i, j]), loadings[i, j], names[j]) for j in range(len(names))),
            reverse=True,
        )
        drivers = [f"{name}({val:+.2f})" for abs_val, val, name in bm_loads if abs_val >= 0.20]
        print(f"PC{i+1:<3} ({var_ratio[i]*100:.1f}%):  " + "  ".join(drivers))

    # ── Mann-Whitney U on retained PCs ───────────────────────────────────────
    treat_idx = groups[groups == 1].index
    ctrl_idx  = groups[groups == 0].index
    # Align scores DataFrame with group indices
    scores_df = pd.DataFrame(scores, index=X_raw.index,
                             columns=[f"PC{i+1}" for i in range(scores.shape[1])])

    print(f"\nGroup Differences on Retained PCs  (Mann-Whitney U)")
    print("=" * 55)
    print(f"{'PC':<6} {'Treat mean':>11} {'Ctrl mean':>10} {'U':>10} {'p':>8}")
    print("-" * 55)

    pc_labels = [f"PC{i+1}" for i in range(n_retain)]
    rows: list[tuple] = []
    for pc in pc_labels:
        t_vals = scores_df.loc[treat_idx, pc].dropna()
        c_vals = scores_df.loc[ctrl_idx,  pc].dropna()
        u_stat, p = mannwhitneyu(t_vals, c_vals, alternative='two-sided')
        rows.append((pc, t_vals.mean(), c_vals.mean(), u_stat, p))

    for pc, t_mean, c_mean, u_stat, p in rows:
        sig = "*" if p < 0.05 else ""
        print(f"{pc:<6} {t_mean:>11.3f} {c_mean:>10.3f} {u_stat:>10.1f} {p:>8.4f}  {sig}")
    print("* p < 0.05")

    _write_docx(rows, n_retain, var_ratio)

    # ── Permutation test: multivariate group separation ───────────────────────
    n_perm    = 10_000
    ret_scores = scores_df[pc_labels].values  # retained PC subspace
    labels     = groups.values

    t_mask = labels == 1
    c_mask = labels == 0
    centroid_t   = ret_scores[t_mask].mean(axis=0)
    centroid_c   = ret_scores[c_mask].mean(axis=0)
    observed_dist = float(np.linalg.norm(centroid_t - centroid_c))

    rng      = np.random.default_rng(42)
    null_dist: list[float] = []
    for _ in range(n_perm):
        perm = rng.permutation(labels)
        ct   = ret_scores[perm == 1].mean(axis=0)
        cc   = ret_scores[perm == 0].mean(axis=0)
        null_dist.append(float(np.linalg.norm(ct - cc)))

    perm_p = float(np.mean(np.array(null_dist) >= observed_dist))

    print(f"\nPermutation Test: Multivariate Group Separation ({n_perm:,} permutations)")
    print("=" * 58)
    print(f"  PC subspace retained:       PC1–PC{n_retain}")
    print(f"  Observed centroid distance: {observed_dist:.4f}")
    print(f"  Null mean ± SD:             {np.mean(null_dist):.4f} ± {np.std(null_dist):.4f}")
    print(f"  Permutation p-value:        {perm_p:.4f}")
    if perm_p < 0.05:
        print("  → Significant multivariate separation (p < 0.05)")
    else:
        print("  → No significant multivariate separation (p ≥ 0.05)")

    # ── Figures ───────────────────────────────────────────────────────────────
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    _plot_scree(eigenvalues, n_kaiser, n_80pct)
    _plot_biplot(scores, loadings, names, groups, var_ratio)
    _plot_centroids(scores, groups, var_ratio, observed_dist, perm_p)
    print(f"\nPlots saved to {_OUT_DIR}/")


def _write_docx(rows: list[tuple], n_retain: int, var_ratio: np.ndarray) -> None:
    """Save Group Differences on Retained PCs table → output/a2_pca.docx."""
    out = Path('output/a2_pca.docx')
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(2.0)
    normal = cast(_ParagraphStyle, doc.styles['Normal'])
    normal.font.name = 'Arial'
    normal.font.size = Pt(FONT_BODY)

    headers = ['PC', 'Variance', 'Treat mean', 'Ctrl mean', 'U statistic', 'p', 'Sig.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_col_widths(table, [1.4, 1.6, 2.2, 2.2, 2.4, 2.0, 1.2])

    for cell, text in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        add_run(cell.paragraphs[0], text, FONT_BODY, bold=True)

    for pc, t_mean, c_mean, u_stat, p in rows:
        sig = '*' if p < 0.05 else ''
        pc_idx = int(pc[2:]) - 1  # PC1 → index 0
        var_pct = f'{var_ratio[pc_idx]*100:.1f}%'
        data = [pc, var_pct, f'{t_mean:.3f}', f'{c_mean:.3f}', f'{u_stat:.1f}', f'{p:.4f}', sig]
        row = table.add_row()
        for cell, val in zip(row.cells, data):
            clear_cell(cell)
            add_run(cell.paragraphs[0], val, FONT_BODY)

    legend = (
        f'Mann-Whitney U test (two-sided), Treatment vs Control, PC1–PC{n_retain} '
        f'(orthogonal components; no multiple-testing correction applied). * p < 0.05.'
    )
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    add_run(para, legend, FONT_LEGEND, rgb=C_GRAY)

    doc.save(out)
    print(f'Saved {out}')


def _plot_scree(eigenvalues: np.ndarray, n_kaiser: int, n_80pct: int) -> None:
    """Scree plot: eigenvalue per PC with Kaiser threshold and 80% cutoff marked."""
    n = len(eigenvalues)
    xs = np.arange(1, n + 1)

    fig, ax1 = plt.subplots(figsize=(8, 5))
    ax2 = ax1.twinx()  # right axis for cumulative variance

    # Bar: eigenvalues; colour by whether above Kaiser threshold
    colors = [_COLOR_TREAT if ev > 1 else '#AAAAAA' for ev in eigenvalues]
    ax1.bar(xs, eigenvalues, color=colors, alpha=0.8, zorder=2)
    ax1.axhline(1.0, color='black', linewidth=1, linestyle='--', label='Kaiser (eigenvalue = 1)')

    # Line: cumulative variance on right axis
    cum_var = np.cumsum(eigenvalues / eigenvalues.sum())
    ax2.plot(xs, cum_var * 100, color=_COLOR_CTRL, marker='o', markersize=4,
             linewidth=1.5, label='Cumulative variance %')
    ax2.axhline(80, color=_COLOR_CTRL, linewidth=0.8, linestyle=':')
    ax2.axhline(90, color='#009E73', linewidth=0.8, linestyle=':')  # Okabe-Ito green
    ax2.set_ylabel('Cumulative variance (%)', color=_COLOR_CTRL)
    ax2.tick_params(axis='y', labelcolor=_COLOR_CTRL)
    ax2.set_ylim(0, 105)

    n_90pct = int(np.searchsorted(cum_var, 0.90)) + 1

    # Vertical markers for retention cutoffs
    ax1.axvline(n_kaiser + 0.5, color=_COLOR_TREAT, linewidth=1, linestyle=':',
                label=f'Kaiser cutoff (PC{n_kaiser})')
    ax1.axvline(n_80pct + 0.5, color=_COLOR_CTRL, linewidth=1, linestyle='-.',
                label=f'80% variance cutoff (PC{n_80pct})')
    ax1.axvline(n_90pct + 0.5, color='#009E73', linewidth=1, linestyle='-.',
                label=f'90% variance cutoff (PC{n_90pct})')

    ax1.set_xlabel('Principal Component')
    ax1.set_ylabel('Eigenvalue')
    ax1.set_title('Scree Plot – PCA of 20 Biomarker Δ Values')
    ax1.set_xticks(xs)
    ax1.set_xticklabels([f'PC{i}' for i in xs], rotation=45, ha='right', fontsize=7)

    # Combine legends from both axes
    h1, l1 = ax1.get_legend_handles_labels()
    h2, l2 = ax2.get_legend_handles_labels()
    ax1.legend(h1 + h2, l1 + l2, fontsize=7, loc='lower center',
               bbox_to_anchor=(0.5, -0.28), ncol=3)

    fig.savefig(_OUT_DIR / 'scree.png', dpi=150, bbox_inches='tight')
    plt.close(fig)


def _plot_biplot(scores: np.ndarray, loadings: np.ndarray, names: list[str],
                 groups: pd.Series, var_ratio: np.ndarray) -> None:
    """PCA biplot: PC1 vs PC2 scores coloured by group, with loading arrows.

    Uses twin axes so score and loading scales are independent — arrows always
    fit inside the plot regardless of score spread.
    """
    fig, ax_scores = plt.subplots(figsize=(9, 7))

    t_mask = groups.values == 1
    c_mask = groups.values == 0

    ax_scores.scatter(scores[t_mask, 0], scores[t_mask, 1],
                      color=_COLOR_TREAT, label='Treatment', alpha=0.75, s=50, zorder=3)
    ax_scores.scatter(scores[c_mask, 0], scores[c_mask, 1],
                      color=_COLOR_CTRL,  label='Control',   alpha=0.75, s=50, zorder=3)

    ax_scores.axhline(0, color='grey', linewidth=0.5, linestyle='--')
    ax_scores.axvline(0, color='grey', linewidth=0.5, linestyle='--')
    ax_scores.set_xlabel(f'PC1 ({var_ratio[0]*100:.1f}% variance)')
    ax_scores.set_ylabel(f'PC2 ({var_ratio[1]*100:.1f}% variance)')
    ax_scores.set_title('PCA Biplot – PC1 vs PC2 (Biomarker Δ Values)')
    ax_scores.legend(fontsize=8, loc='upper left')

    # Loading arrows on a separate twin axis with its own [-1, 1] scale
    ax_load = ax_scores.twinx().twiny()  # independent x and y scales
    ax_load.set_xlim(-1, 1)
    ax_load.set_ylim(-1, 1)
    ax_load.set_xticks([])
    ax_load.set_yticks([])

    arrow_len_2d = np.sqrt(loadings[0] ** 2 + loadings[1] ** 2)  # length in PC1/PC2 plane
    threshold    = 0.35 * arrow_len_2d.max()

    texts = []
    for j, name in enumerate(names):
        lx, ly = loadings[0, j], loadings[1, j]
        if arrow_len_2d[j] < threshold:
            continue
        ax_load.annotate('', xy=(lx, ly), xytext=(0, 0),
                         arrowprops=dict(arrowstyle='->', color='#CC79A7', lw=1.2))
        # Place label just past the arrowhead; adjust_text will move it if needed
        texts.append(ax_load.text(lx * 1.1, ly * 1.1, name,
                                  fontsize=7, color='#CC79A7', ha='center', va='center'))

    # Nudge overlapping labels off each other and away from arrow lines
    from adjustText import adjust_text
    adjust_text(texts, ax=ax_load, arrowprops=dict(arrowstyle='-', color='#CC79A7', lw=0.5))

    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'biplot.png', dpi=150)
    plt.close(fig)


def _confidence_ellipse(x: np.ndarray, y: np.ndarray, ax: plt.Axes,
                        color: str, n_std: float = 1.96) -> None:
    """Draw a 95% confidence ellipse (1.96 SD) for a cloud of 2-D points."""
    from matplotlib.patches import Ellipse
    import matplotlib.transforms as transforms

    cov = np.cov(x, y)
    vals, vecs = np.linalg.eigh(cov)
    order = vals.argsort()[::-1]
    vals, vecs = vals[order], vecs[:, order]
    angle = np.degrees(np.arctan2(*vecs[:, 0][::-1]))
    width, height = 2 * n_std * np.sqrt(vals)
    ellipse = Ellipse(
        xy=(x.mean(), y.mean()), width=width, height=height, angle=angle,
        edgecolor=color, facecolor=color, alpha=0.12, linewidth=1.5, linestyle='--',
    )
    ax.add_patch(ellipse)


def _plot_centroids(scores: np.ndarray, groups: pd.Series,
                    var_ratio: np.ndarray, obs_dist: float, perm_p: float) -> None:
    """PC1 vs PC2 scatter with group centroids, 95% ellipses, and centroid connector."""
    fig, ax = plt.subplots(figsize=(7, 6))

    t_mask = groups.values == 1
    c_mask = groups.values == 0

    t_scores = scores[t_mask, :2]
    c_scores = scores[c_mask, :2]

    # Individual points (semi-transparent)
    ax.scatter(t_scores[:, 0], t_scores[:, 1],
               color=_COLOR_TREAT, alpha=0.35, s=35, zorder=2)
    ax.scatter(c_scores[:, 0], c_scores[:, 1],
               color=_COLOR_CTRL,  alpha=0.35, s=35, zorder=2)

    # 95% confidence ellipses
    _confidence_ellipse(t_scores[:, 0], t_scores[:, 1], ax, _COLOR_TREAT)
    _confidence_ellipse(c_scores[:, 0], c_scores[:, 1], ax, _COLOR_CTRL)

    # Centroids
    ct = t_scores.mean(axis=0)
    cc = c_scores.mean(axis=0)
    ax.scatter(*ct, color=_COLOR_TREAT, s=120, zorder=4, marker='D', label='Treatment centroid')
    ax.scatter(*cc, color=_COLOR_CTRL,  s=120, zorder=4, marker='D', label='Control centroid')

    # Line connecting centroids, labelled with distance and p-value
    ax.plot([ct[0], cc[0]], [ct[1], cc[1]], color='#555555', linewidth=1.2,
            linestyle='--', zorder=3)
    mid = (ct + cc) / 2
    sig_label = 'p < 0.05' if perm_p < 0.05 else f'p = {perm_p:.3f}'
    ax.annotate(
        f'd = {obs_dist:.3f}\n{sig_label} (permutation)',
        xy=mid, xytext=(mid[0] + 0.3, mid[1] + 0.4),
        fontsize=7.5, color='#555555',
        arrowprops=dict(arrowstyle='->', color='#555555', lw=0.8),
    )

    # Dummy patches for group legend entries (include ellipse in label)
    from matplotlib.patches import Patch
    legend_handles = [
        Patch(facecolor=_COLOR_TREAT, alpha=0.5, label='Treatment (points + 95% ellipse)'),
        Patch(facecolor=_COLOR_CTRL,  alpha=0.5, label='Control (points + 95% ellipse)'),
        plt.Line2D([0], [0], marker='D', color='w', markerfacecolor=_COLOR_TREAT, markersize=8, label='Treatment centroid'),
        plt.Line2D([0], [0], marker='D', color='w', markerfacecolor=_COLOR_CTRL,  markersize=8, label='Control centroid'),
    ]
    ax.legend(handles=legend_handles, fontsize=7.5, loc='upper left')

    ax.axhline(0, color='grey', linewidth=0.5, linestyle='--')
    ax.axvline(0, color='grey', linewidth=0.5, linestyle='--')
    ax.set_xlabel(f'PC1 ({var_ratio[0]*100:.1f}% variance)')
    ax.set_ylabel(f'PC2 ({var_ratio[1]*100:.1f}% variance)')
    ax.set_title('Group Separation in PC Space (PC1 vs PC2)')

    fig.tight_layout()
    fig.savefig(_OUT_DIR / 'centroids.png', dpi=150)
    plt.close(fig)
