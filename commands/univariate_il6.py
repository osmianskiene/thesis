"""Two-part IL-6 analysis addressing left-censoring at the assay LOD (2.0 pg/mL).

Part 1 — logistic regression: P(IL-6 detectable at T1) ~ Treatment + I(detectable at T0)
Part 2 — OLS on log scale: log(IL-6 T1) ~ Treatment + log(IL-6 T0), restricted to T1-detectable
Sensitivity — LOD/2 substitution: Mann-Whitney U on Δ with values-at-LOD replaced by LOD/2
"""
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd
import scipy.stats as stats
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt, Cm
from docx.styles.style import _ParagraphStyle

from data import load_composites
from docx_utils import (
    C_BLACK, C_GRAY, FONT_BODY, FONT_LEGEND,
    add_run, clear_cell, new_para, set_col_widths,
)

IL6_LOD = 2.0  # assay lower limit of detection in pg/mL

_OUT = Path('output/a1_univariate_il6.docx')


@dataclass
class _CensorSummary:
    """Censoring counts at T0 and T1 by group."""
    n_total: int
    n_lod_t0: int
    n_lod_t1: int
    n_lod_t0_treat: int
    n_lod_t0_ctrl: int
    n_lod_t1_treat: int
    n_lod_t1_ctrl: int
    n_treat: int
    n_ctrl: int


@dataclass
class _Part1Result:
    """Logistic regression result for P(detect at T1)."""
    n: int
    or_treat: float
    ci_lo_treat: float
    ci_hi_treat: float
    p_treat: float
    or_detect0: float
    ci_lo_detect0: float
    ci_hi_detect0: float
    p_detect0: float
    pseudo_r2: float
    n_det_treat: int
    n_treat: int
    n_det_ctrl: int
    n_ctrl: int


@dataclass
class _Part2Result:
    """OLS result on log(IL-6 T1) among detectable T1 values."""
    n: int
    n_treat: int
    n_ctrl: int
    beta_treat: float
    ci_lo: float
    ci_hi: float
    p: float
    ratio: float  # exp(beta_treat)
    r2: float


@dataclass
class _SensResult:
    """LOD/2 substitution sensitivity result."""
    mean_delta_treat: float
    median_delta_treat: float
    mean_delta_ctrl: float
    median_delta_ctrl: float
    p_mw: float
    cohens_d: float


def _compute_censor_summary(
    il6_t0: pd.Series, il6_t1: pd.Series, group: pd.Series
) -> _CensorSummary:
    n_total = len(il6_t0)
    at_lod_t0 = il6_t0 == IL6_LOD
    at_lod_t1 = il6_t1 == IL6_LOD
    return _CensorSummary(
        n_total=n_total,
        n_lod_t0=int(at_lod_t0.sum()),
        n_lod_t1=int(at_lod_t1.sum()),
        n_lod_t0_treat=int((at_lod_t0 & (group == 1)).sum()),
        n_lod_t0_ctrl=int( (at_lod_t0 & (group == 0)).sum()),
        n_lod_t1_treat=int((at_lod_t1 & (group == 1)).sum()),
        n_lod_t1_ctrl=int( (at_lod_t1 & (group == 0)).sum()),
        n_treat=int((group == 1).sum()),
        n_ctrl=int((group == 0).sum()),
    )


def _fit_part1(
    il6_t0: pd.Series, il6_t1: pd.Series, group: pd.Series
) -> _Part1Result | None:
    df = pd.DataFrame({
        'detect_t1': (il6_t1 > IL6_LOD).astype(int),
        'detect_t0': (il6_t0 > IL6_LOD).astype(int),
        'Treatment': group,
    }).dropna()

    X = sm.add_constant(df[['Treatment', 'detect_t0']])
    try:
        model = sm.Logit(df['detect_t1'], X).fit(disp=False)
    except Exception:
        return None

    def _ci(var: str) -> tuple[float, float]:
        c = model.params[var]
        s = model.bse[var]
        return float(np.exp(c - 1.96 * s)), float(np.exp(c + 1.96 * s))

    ci_t  = _ci('Treatment')
    ci_d0 = _ci('detect_t0')

    n_det_treat = int(df.loc[df['Treatment'] == 1, 'detect_t1'].sum())
    n_treat     = int((df['Treatment'] == 1).sum())
    n_det_ctrl  = int(df.loc[df['Treatment'] == 0, 'detect_t1'].sum())
    n_ctrl      = int((df['Treatment'] == 0).sum())

    return _Part1Result(
        n=len(df),
        or_treat=float(np.exp(model.params['Treatment'])),
        ci_lo_treat=ci_t[0],
        ci_hi_treat=ci_t[1],
        p_treat=float(model.pvalues['Treatment']),
        or_detect0=float(np.exp(model.params['detect_t0'])),
        ci_lo_detect0=ci_d0[0],
        ci_hi_detect0=ci_d0[1],
        p_detect0=float(model.pvalues['detect_t0']),
        pseudo_r2=float(model.prsquared),
        n_det_treat=n_det_treat,
        n_treat=n_treat,
        n_det_ctrl=n_det_ctrl,
        n_ctrl=n_ctrl,
    )


def _fit_part2(
    il6_t0: pd.Series, il6_t1: pd.Series, group: pd.Series
) -> _Part2Result | None:
    # Restrict to participants with detectable IL-6 at T1
    mask = (il6_t1 > IL6_LOD) & (il6_t0 > 0)
    df = pd.DataFrame({
        'log_il6_t1': np.log(il6_t1[mask]),
        # Floor T0 at LOD before log to avoid log(0) for T0-censored participants
        'log_il6_t0': np.log(il6_t0[mask].clip(lower=IL6_LOD)),
        'Treatment': group[mask],
    }).dropna()

    if len(df) < 10:
        return None

    X = sm.add_constant(df[['Treatment', 'log_il6_t0']])
    model = sm.OLS(df['log_il6_t1'], X).fit()

    beta  = float(model.params['Treatment'])
    ci    = model.conf_int().loc['Treatment']
    return _Part2Result(
        n=len(df),
        n_treat=int((df['Treatment'] == 1).sum()),
        n_ctrl=int((df['Treatment'] == 0).sum()),
        beta_treat=beta,
        ci_lo=float(ci[0]),
        ci_hi=float(ci[1]),
        p=float(model.pvalues['Treatment']),
        ratio=float(np.exp(beta)),
        r2=float(model.rsquared),
    )


def _fit_sensitivity(
    il6_t0: pd.Series, il6_t1: pd.Series, group: pd.Series
) -> _SensResult:
    # Replace values at LOD with LOD/2 (standard left-censoring substitution)
    t0_sub = il6_t0.where(il6_t0 > IL6_LOD, IL6_LOD / 2)
    t1_sub = il6_t1.where(il6_t1 > IL6_LOD, IL6_LOD / 2)
    delta  = t1_sub - t0_sub

    t_delta = delta[group == 1].dropna()
    c_delta = delta[group == 0].dropna()

    _, p_mw = stats.mannwhitneyu(t_delta, c_delta, alternative='two-sided')

    # Pooled Cohen's d on substituted Δ
    na, nb = len(t_delta), len(c_delta)
    pooled_sd = np.sqrt(
        ((na - 1) * t_delta.std(ddof=1) ** 2 + (nb - 1) * c_delta.std(ddof=1) ** 2)
        / (na + nb - 2)
    )
    d = float((t_delta.mean() - c_delta.mean()) / pooled_sd) if pooled_sd > 0 else float('nan')

    return _SensResult(
        mean_delta_treat=float(t_delta.mean()),
        median_delta_treat=float(t_delta.median()),
        mean_delta_ctrl=float(c_delta.mean()),
        median_delta_ctrl=float(c_delta.median()),
        p_mw=float(p_mw),
        cohens_d=d,
    )


# ---------------------------------------------------------------------------
# stdout output
# ---------------------------------------------------------------------------

def _print_results(
    cs: _CensorSummary,
    p1: _Part1Result | None,
    p2: _Part2Result | None,
    sens: _SensResult,
) -> None:
    pct = lambda n, d: f'{100 * n / d:.1f}%'  # noqa: E731  # inline helper

    print('IL-6 TWO-PART CENSORED ANALYSIS')
    print(f'LOD = {IL6_LOD} pg/mL')
    print()
    print('Censoring summary')
    print(f'  T0 at LOD: {cs.n_lod_t0}/{cs.n_total} ({pct(cs.n_lod_t0, cs.n_total)})'
          f'  [Treatment {cs.n_lod_t0_treat}/{cs.n_treat}, Control {cs.n_lod_t0_ctrl}/{cs.n_ctrl}]')
    print(f'  T1 at LOD: {cs.n_lod_t1}/{cs.n_total} ({pct(cs.n_lod_t1, cs.n_total)})'
          f'  [Treatment {cs.n_lod_t1_treat}/{cs.n_treat}, Control {cs.n_lod_t1_ctrl}/{cs.n_ctrl}]')

    print()
    print('Part 1 — Logistic: P(IL-6 detectable at T1) ~ Treatment + I(detectable at T0)')
    if p1 is None:
        print('  Model failed to converge.')
    else:
        print(f'  N = {p1.n}')
        print(f'  {"Variable":<20} {"OR":>6}  {"95% CI":>16}  {"p":>7}')
        print(f'  {"-"*53}')
        for lbl, or_, lo, hi, p in [
            ('Treatment',       p1.or_treat,    p1.ci_lo_treat,    p1.ci_hi_treat,    p1.p_treat),
            ('Detectable at T0', p1.or_detect0, p1.ci_lo_detect0, p1.ci_hi_detect0, p1.p_detect0),
        ]:
            sig = '*' if p < 0.05 else ''
            print(f'  {lbl:<20} {or_:>6.3f}  [{lo:.3f}, {hi:.3f}]  {p:>7.4f} {sig}')
        print(f'  Pseudo R² (McFadden) = {p1.pseudo_r2:.3f}')
        print(f'  Detectable at T1: Treatment {p1.n_det_treat}/{p1.n_treat}'
              f' ({pct(p1.n_det_treat, p1.n_treat)}), '
              f'Control {p1.n_det_ctrl}/{p1.n_ctrl}'
              f' ({pct(p1.n_det_ctrl, p1.n_ctrl)})')

    print()
    print('Part 2 — OLS: log(IL-6 T1) ~ Treatment + log(IL-6 T0), T1-detectable only')
    if p2 is None:
        print('  Insufficient detectable T1 values for regression (n < 10).')
    else:
        sig = '*' if p2.p < 0.05 else ''
        print(f'  N = {p2.n} (Treatment {p2.n_treat}, Control {p2.n_ctrl})')
        print(f'  Treatment β = {p2.beta_treat:.3f}  (95% CI [{p2.ci_lo:.3f}, {p2.ci_hi:.3f}])')
        print(f'  exp(β) = {p2.ratio:.3f}×  (IL-6 ratio treatment vs control)')
        print(f'  p = {p2.p:.4f}  R² = {p2.r2:.3f}  {sig}')

    print()
    print(f'Sensitivity — LOD/2 substitution (censored values replaced with {IL6_LOD / 2} pg/mL)')
    print(f'  Treatment Δ: mean={sens.mean_delta_treat:.3f}, median={sens.median_delta_treat:.3f}')
    print(f'  Control   Δ: mean={sens.mean_delta_ctrl:.3f}, median={sens.median_delta_ctrl:.3f}')
    sig = '*' if sens.p_mw < 0.05 else ''
    print(f'  Mann-Whitney U  p = {sens.p_mw:.4f}  Cohen\'s d = {sens.cohens_d:.3f}  {sig}')
    print('  Note: LOD/2 substitution biases toward the null; use as cross-check only.')


# ---------------------------------------------------------------------------
# docx output
# ---------------------------------------------------------------------------

def _body_style(doc: Document) -> '_ParagraphStyle':
    style = cast = doc.styles['Normal']
    return style  # type: ignore[return-value]


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, FONT_BODY, bold=True)


def _note_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, FONT_LEGEND, rgb=C_GRAY)


def _build_docx(
    cs: _CensorSummary,
    p1: _Part1Result | None,
    p2: _Part2Result | None,
    sens: _SensResult,
) -> None:
    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    add_run(title, 'IL-6 Two-Part Censored Analysis', FONT_BODY + 2, bold=True)

    pct = lambda n, d: f'{100 * n / d:.1f}%'  # noqa: E731

    # --- Censoring summary ---
    _section_heading(doc, 'Censoring summary')
    tbl_cs = doc.add_table(rows=3, cols=4)
    tbl_cs.style = 'Table Grid'
    set_col_widths(tbl_cs, [3.5, 2.5, 2.5, 2.5])

    headers = ['', 'Total', 'Treatment', 'Control']
    for col_i, h in enumerate(headers):
        p = tbl_cs.rows[0].cells[col_i].paragraphs[0]
        add_run(p, h, FONT_BODY, bold=True)

    rows_cs = [
        ('At LOD at T0',
         f'{cs.n_lod_t0}/{cs.n_total} ({pct(cs.n_lod_t0, cs.n_total)})',
         f'{cs.n_lod_t0_treat}/{cs.n_treat} ({pct(cs.n_lod_t0_treat, cs.n_treat)})',
         f'{cs.n_lod_t0_ctrl}/{cs.n_ctrl} ({pct(cs.n_lod_t0_ctrl, cs.n_ctrl)})'),
        ('At LOD at T1',
         f'{cs.n_lod_t1}/{cs.n_total} ({pct(cs.n_lod_t1, cs.n_total)})',
         f'{cs.n_lod_t1_treat}/{cs.n_treat} ({pct(cs.n_lod_t1_treat, cs.n_treat)})',
         f'{cs.n_lod_t1_ctrl}/{cs.n_ctrl} ({pct(cs.n_lod_t1_ctrl, cs.n_ctrl)})'),
    ]
    for row_i, row_data in enumerate(rows_cs):
        for col_i, text in enumerate(row_data):
            p = tbl_cs.rows[row_i + 1].cells[col_i].paragraphs[0]
            add_run(p, text, FONT_BODY)

    _note_para(doc, f'LOD = {IL6_LOD} pg/mL (assay lower limit of detection).')

    # --- Part 1: logistic ---
    _section_heading(doc, 'Part 1 — Logistic regression: P(IL-6 detectable at T1)')
    _note_para(doc, 'Model: I(IL-6 T1 > LOD) ~ Treatment + I(IL-6 T0 > LOD)')

    if p1 is None:
        p = doc.add_paragraph()
        add_run(p, 'Model failed to converge.', FONT_BODY, rgb=C_GRAY)
    else:
        tbl_p1 = doc.add_table(rows=3, cols=4)
        tbl_p1.style = 'Table Grid'
        set_col_widths(tbl_p1, [4.0, 2.0, 3.0, 2.0])

        for col_i, h in enumerate(['Variable', 'OR', '95% CI', 'p']):
            p = tbl_p1.rows[0].cells[col_i].paragraphs[0]
            add_run(p, h, FONT_BODY, bold=True)

        p1_rows = [
            ('Treatment',
             f'{p1.or_treat:.3f}',
             f'[{p1.ci_lo_treat:.3f}, {p1.ci_hi_treat:.3f}]',
             f'{p1.p_treat:.4f}{"*" if p1.p_treat < 0.05 else ""}'),
            ('Detectable at T0',
             f'{p1.or_detect0:.3f}',
             f'[{p1.ci_lo_detect0:.3f}, {p1.ci_hi_detect0:.3f}]',
             f'{p1.p_detect0:.4f}{"*" if p1.p_detect0 < 0.05 else ""}'),
        ]
        for row_i, row_data in enumerate(p1_rows):
            for col_i, text in enumerate(row_data):
                p_cell = tbl_p1.rows[row_i + 1].cells[col_i].paragraphs[0]
                add_run(p_cell, text, FONT_BODY)

        _note_para(doc,
            f'N = {p1.n}   Pseudo R² (McFadden) = {p1.pseudo_r2:.3f}\n'
            f'Detectable at T1: Treatment {p1.n_det_treat}/{p1.n_treat}'
            f' ({pct(p1.n_det_treat, p1.n_treat)}), '
            f'Control {p1.n_det_ctrl}/{p1.n_ctrl}'
            f' ({pct(p1.n_det_ctrl, p1.n_ctrl)})'
        )

    # --- Part 2: OLS ---
    _section_heading(doc, 'Part 2 — OLS: log(IL-6 T1) among detectable values')
    _note_para(doc, 'Model: log(IL-6 T1) ~ Treatment + log(IL-6 T0)   [T1 > LOD only]')

    if p2 is None:
        p = doc.add_paragraph()
        add_run(p, 'Insufficient detectable T1 values for regression (n < 10).', FONT_BODY, rgb=C_GRAY)
    else:
        tbl_p2 = doc.add_table(rows=2, cols=5)
        tbl_p2.style = 'Table Grid'
        set_col_widths(tbl_p2, [1.5, 2.0, 3.0, 1.5, 1.5])

        for col_i, h in enumerate(['N', 'β (Treatment)', '95% CI', 'p', 'R²']):
            p = tbl_p2.rows[0].cells[col_i].paragraphs[0]
            add_run(p, h, FONT_BODY, bold=True)

        p2_vals = [
            str(p2.n),
            f'{p2.beta_treat:.3f}',
            f'[{p2.ci_lo:.3f}, {p2.ci_hi:.3f}]',
            f'{p2.p:.4f}{"*" if p2.p < 0.05 else ""}',
            f'{p2.r2:.3f}',
        ]
        for col_i, text in enumerate(p2_vals):
            p_cell = tbl_p2.rows[1].cells[col_i].paragraphs[0]
            add_run(p_cell, text, FONT_BODY)

        _note_para(doc,
            f'exp(β) = {p2.ratio:.3f}× (IL-6 ratio treatment vs control among detectable values).  '
            f'N: Treatment {p2.n_treat}, Control {p2.n_ctrl}.'
        )

    # --- Sensitivity ---
    _section_heading(doc, f'Sensitivity — LOD/2 substitution ({IL6_LOD / 2} pg/mL)')
    _note_para(doc, 'Values at LOD replaced with LOD/2; Mann-Whitney U on Δ(T1−T0).')

    tbl_s = doc.add_table(rows=3, cols=3)
    tbl_s.style = 'Table Grid'
    set_col_widths(tbl_s, [3.5, 2.5, 2.5])

    for col_i, h in enumerate(['', 'Mean Δ', 'Median Δ']):
        p = tbl_s.rows[0].cells[col_i].paragraphs[0]
        add_run(p, h, FONT_BODY, bold=True)

    for row_i, (lbl, mn, med) in enumerate([
        ('Treatment', sens.mean_delta_treat, sens.median_delta_treat),
        ('Control',   sens.mean_delta_ctrl,  sens.median_delta_ctrl),
    ]):
        for col_i, text in enumerate([lbl, f'{mn:.3f}', f'{med:.3f}']):
            p_cell = tbl_s.rows[row_i + 1].cells[col_i].paragraphs[0]
            add_run(p_cell, text, FONT_BODY)

    sig = '*' if sens.p_mw < 0.05 else ''
    _note_para(doc,
        f'Mann-Whitney U  p = {sens.p_mw:.4f}{sig}   Cohen\'s d = {sens.cohens_d:.3f}\n'
        'Caution: LOD/2 substitution biases toward the null.'
    )

    _OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(_OUT)
    print(f'\nSaved → {_OUT}')


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def univariate_il6() -> None:
    """Two-part IL-6 analysis (left-censoring) → stdout + output/a1_univariate_il6.docx"""
    df   = load_composites()
    il6_t0 = df['IL 6 T0'].copy()
    il6_t1 = df['IL 6 T1'].copy()
    group  = df['Treatment'].copy()

    cs   = _compute_censor_summary(il6_t0, il6_t1, group)
    p1   = _fit_part1(il6_t0, il6_t1, group)
    p2   = _fit_part2(il6_t0, il6_t1, group)
    sens = _fit_sensitivity(il6_t0, il6_t1, group)

    _print_results(cs, p1, p2, sens)
    _build_docx(cs, p1, p2, sens)
